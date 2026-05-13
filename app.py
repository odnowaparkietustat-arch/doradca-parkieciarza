import streamlit as st
from datetime import date
import io
import math
import io

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import fpdf
    from fpdf import FPDF
    EXPORTS_READY = True
except ImportError:
    EXPORTS_READY = False

# ==========================================
# 1. KONFIGURACJA STRONY I WSPÓLNE FUNKCJE
# ==========================================
st.set_page_config(page_title="Ekspert Parkieciarski WAKOL", layout="wide")

class ReportBuilder:
    def __init__(self):
        self.md_lines = []
    
    def write(self, text):
        self.md_lines.append(str(text))
        
    def markdown(self, text):
        self.md_lines.append(str(text))
        
    def error(self, text):
        st.error(text)
        
    def get_markdown(self):
        return "\n\n".join(self.md_lines)

# --- STAŁE TECHNOLOGICZNE (OPISY PRODUKTÓW) ---
FULL_PS275 = "* Zalecamy aplikację gruntówki wzmacniającej **WAKOL PS 275** w dwóch warstwach – grubym wałkiem sznurkowym, zużycie w sumie **ok. 700 g/m²**. Każda z warstw po **350 g/m²**, aplikowane po sobie w odstępie jednej godziny. Aplikując gruntówkę **WAKOL PS 275** należy zwrócić uwagę, aby dobrze wchłaniała się w podłoże i unikać powstawania kałuż na powierzchni jastrychu. Po nałożeniu drugiej warstwy gruntówki w razie potrzeby wykonać posypkę z piasku kwarcowego. **Po 7 dniach schnięcia** powierzchnię należy **przeszlifować papierem o gradacji 24 – 40** usuwając przyklejony do powierzchni piasek kwarcowy i dokładnie odkurzyć."
FULL_PU235_1W = "* Zalecamy wykonanie gruntowania wzmacniającego poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **WAKOL PU 235**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie **ok. 150 g/m²**. **Czas schnięcia – 4-6 godzin**."
FULL_PU235_BARRIER = "* Zalecamy wykonanie **bariery przeciwwilgociowej** poprzez dwukrotne zagruntowanie gruntówką wzmacniającą **WAKOL PU 235**. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **ok. 150 g/m²**. **Czas schnięcia – 4-6 godzin**. 2. warstwa zużycie **ok. 100 g/m²**. **Czas schnięcia – 4-6 godzin**. **Czas klejenia 72 godziny od zagruntowania**."
FULL_PU280_1W = "* Zalecamy wykonanie gruntowania wzmacniającego poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **{PRODUCTS['PU 280 (1W)']['name']}**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie **ok. 150 g/m²**. **Czas schnięcia – jedna godzina**."
FULL_PU280_BARRIER = "* Z uwagi na podwyższoną wilgotność zalecamy stworzenie **bariery przeciwwilgociowej** poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **{PRODUCTS['PU 280 (1W)']['name']}**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **150 g/m²**. **Czas schnięcia – jedna godzina**. 2. warstwa **100 g/m²** - **czas schnięcia – jedna godzina**. **Czas do klejenia: 72 godziny od zagruntowania**."
FULL_PU280_BARRIER_PLYTA = "* Z uwagi na grubość płyty fundamentowej zalecamy stworzenie **bariery przeciwwilgociowej** poprzez zagruntowanie powierzchni podłoża gruntówką poliuretanową **{PRODUCTS['PU 280 (1W)']['name']}**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **150 g/m²**. **Czas schnięcia – jedna godzina**. 2. warstwa **100 g/m²** - **czas schnięcia – jedna godzina**. **Czas do klejenia: 72 godziny od zagruntowania**."
FULL_D3004 = "* Zagruntować podłoże koncentratem gruntówki dyspersyjnej **WAKOL D 3004**. Proporcje mieszania: 1 część **WAKOL D 3004** + 2 części wody. **Czas schnięcia**: na jastrychach cementowych i betonie po optycznym wyschnięciu **ok. 30 min**. Sposób nanoszenia: wałek do gruntowania microfazer. Zużycie: **ok. 50 g/m²** koncentratu."
FULL_Z625 = "* Wylać masę wyrównawczą **WAKOL Z 625** - wymieszać ją w czystym naczyniu z zimną wodą w proporcji 6,00 – 6,25 litrów wody na 25 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,6 kg/m²/mm**. **Możliwość chodzenia po 2 godzinach**. **Możliwość klejenia podłóg drewnianych przy warstwie do 5 mm – po 6 godzinach**, przy warstwie do 10 mm – po 12 godzinach, przy warstwie 30 mm – po 24 godzinach."
FULL_Z675 = "* Wylać masę wyrównawczą **WAKOL Z 675** - wymieszać ją w czystym naczyniu z zimną wodą w proporcji 6,0 – 6,5 litrów wody na 25 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,6 kg/m²/mm**. **Możliwość chodzenia po 2-3 godzinach**. **Możliwość klejenia podłóg po ok. 24 godzinach przy grubości warstwy do 3 mm**, przy większych grubościach czas schnięcia ulega wydłużeniu."
FULL_Z635 = "* Następnie na podłoże wylać masę wyrównawczą **WAKOL Z 635** - wymieszać ją w czystym naczyniu z zimną wodą w proporcji 6,25 litrów wody na 25 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,6 kg/m²/mm**. **Możliwość chodzenia po 2,5 godzinach**. **Możliwość klejenia podłóg drewnianych przy warstwie do 5 mm – po 24 godzinach**, przy warstwie do 10 mm – po 72 godzinach."
FULL_D3055 = "* Zalecamy zagruntowanie całej powierzchni jastrychu gruntówką dyspersyjną **WAKOL D 3055** - aplikacja wałkiem **ok. 150 g/m²**. **Czas schnięcia ok. 30 min**."

def insert_header():
    logo_wakol = "https://www.wakol.com/fileadmin/templates/images/wakol_logo.png"
    st.markdown(f"""
    <div style="display: flex; justify-content: space-between; align-items: center; background-color: #f8f9fa; padding: 20px; border-radius: 10px; border: 1px solid #dee2e6;">
        <div style="flex: 1;"><img src="{logo_wakol}" width="220">
            <div style="margin-top: 15px; font-size: 11px; color: #555; line-height: 1.4;">
                <b>Loba-Wakol Polska Sp. z o.o.</b><br>ul. Sławęcińska 16, Macierzysz | 05-850 Ożarów Mazowiecki<br>tel.: +48 22 436 24 20 | biuro@loba-wakol.pl
            </div>
        </div>
        <div style="flex: 1; text-align: right;">
            <div style="font-size: 18px; font-weight: bold; color: #000; margin-bottom: 5px;">PROTOKÓŁ TECHNICZNY</div>
            <div style="font-size: 14px; color: #333;">Anspruch verbindet</div>
        </div>
    </div><br>
    """, unsafe_allow_html=True)

# ==========================================
# 2. LOGIKA DLA POSZCZEGÓLNYCH OKŁADZIN
# ==========================================

def render_wspolne_dane_optyczne(dane, rep):
    age_txt = f" w wieku {dane['substrate_age_val']} miesięcy" if dane['substrate_age_val'] else ""
    heat_txt = f" Została zainstalowana {dane['heating_info']}." if dane['heating_exists'] == "TAK" else " Brak instalacji ogrzewania podłogowego."
    curing_txt = " Został przeprowadzony proces wygrzewania zgodnie z protokołem." if dane['heating_curing_done'] == "TAK" else " Nie został przeprowadzony proces wygrzewania podłoża." if dane['heating_exists'] == "TAK" else ""
    dil_txt = " Dylatacje zachowane prawidłowo." if dane['dilatations_obw_ok'] == "TAK" else " **Dylatacje obwodowe nie zachowane prawidłowo.**"
    klaw_m = dane.get('klaw_meters') or 0
    pek_m = dane.get('pek_meters') or 0
    klaw_txt = f" **Zaobserwowano {klaw_m} metrów bieżących dylatacji pozornych wymagających zespolenia.**" if dane['cracks_klaw'] == "TAK" else ""
    pek_txt = f" **Stwierdzono obecność pęknięć wymagających zespolenia ({pek_m} mb).**" if dane['cracks_pek'] == "TAK" else ""
    holes_txt = f" **Zlokalizowano ubytki wymagające wypełnienia masą naprawczą{dane['hole_details']}.**" if dane['holes'] == "TAK" else ""
    demolition_txt = " **Podłoże wymaga demontażu przed przystąpieniem do dalszych prac.**" if dane.get('requires_demolition') else ""
    level_txt = f" **Podłoże wymaga wyrównania masą wyrównawczą o planowanej grubości {dane['leveling_thickness']} milimetrów.**" if dane['needs_levelling'] == "TAK" else ""
    vent_txt = f" Rodzaj zastosowanej wentylacji: wentylacja {dane['ventilation_type'].lower()}."
    evenness_txt = " Nie badano równości podłoża." if dane['needs_levelling'] == "NIE" else ""
    dodatkowe_txt = f" **{dane['dodatkowe_informacje']}**" if dane.get('dodatkowe_informacje') else ""

    area_txt = f" o powierzchni {dane['area_m2']} m²" if dane.get('area_m2') else ""
    full_opt_report = f"Podłoże pod planowaną okładzinę ({dane['flooring_type']}) stanowi {dane['substrate']}{area_txt}{age_txt}.{dodatkowe_txt}{demolition_txt}{heat_txt}{curing_txt}{dil_txt}{klaw_txt}{pek_txt}{holes_txt}{level_txt} {vent_txt}{evenness_txt}"
    rep.write(f"**a) oględziny optyczne:** {full_opt_report}")
    
    presso_valid = [str(p) for p in dane.get('presso_results', []) if p is not None]
    
    tests_out = []
    if dane.get('test_hammer'): tests_out.append(f"- Młotek: {dane['test_hammer']}")
    if dane.get('test_ripper'): tests_out.append(f"- Rysik: {dane['test_ripper']}")
    if dane.get('test_brush'): tests_out.append(f"- Szczotka: {dane['test_brush']}")
    if presso_valid: tests_out.append(f"- Wyniki PressoMess: {', '.join(presso_valid)} N/mm²")
    tests_out.append(f"- Ocena ogólna: **{dane['strength_labels'][dane['strength_val']]}**")
    
    tests_str = "\n".join(tests_out)
    rep.write(f"**b) badanie wytrzymałości:**\n{tests_str}")
    
    if dane.get('moisture') is not None:
        moisture_status = "POZYTYWNY" if dane['moisture'] <= dane['limit'] else "NEGATYWNY"
        rep.write(f"**c) badanie wilgotności:** Wynik badania wilgotności metodą CM: **{dane['moisture']} % CM** (Norma: {dane['limit']} % CM) — **Wynik: {moisture_status}**")
    else:
        rep.write("**c) badanie wilgotności:** Nie dotyczy — podłoże drewniane / ceramiczne.")

    klimat = []
    if dane.get('temp_air') is not None: klimat.append(f"Temperatura powietrza: {dane['temp_air']} °C")
    if dane.get('hum_air') is not None: klimat.append(f"Wilgotność powietrza: {dane['hum_air']} %")
    if klimat:
        rep.write(f"**d) warunki klimatyczne:** {', '.join(klimat)}.")

PRODUCTS_WAKOL = {
    'PU 280 (1W)': {'name': 'WAKOL PU 280 (1 warstwa)', 'usage': 150, 'sizes': [11, 5], 'text': FULL_PU280_1W, 'price': 54.27},
    'PU 280 (Bariera)': {'name': 'WAKOL PU 280 (bariera)', 'usage': 250, 'sizes': [11, 5], 'text': FULL_PU280_BARRIER, 'price': 54.27},
    'PU 280 (Bariera Płyta)': {'name': 'WAKOL PU 280 (bariera)', 'usage': 250, 'sizes': [11, 5], 'text': FULL_PU280_BARRIER_PLYTA, 'price': 54.27},
    'PU 235 (1W)': {'name': 'WAKOL PU 235 (1 warstwa)', 'usage': 150, 'sizes': [11], 'text': FULL_PU235_1W, 'price': 50.60},
    'PU 235 (Bariera)': {'name': 'WAKOL PU 235 (bariera)', 'usage': 250, 'sizes': [11], 'text': FULL_PU235_BARRIER, 'price': 50.60},
    'PS 275': {'name': 'WAKOL PS 275', 'usage': 700, 'sizes': [11], 'text': FULL_PS275, 'price': 19.08},
    'D 3004': {'name': 'WAKOL D 3004', 'usage': 50, 'sizes': [10, 5], 'text': FULL_D3004, 'price': 23.32},
    'D 3045': {'name': 'WAKOL D 3045 (mostek sczepny)', 'usage': 150, 'sizes': [12, 6], 'text': "", 'price': 25.00},
    'D 3055': {'name': 'WAKOL D 3055', 'usage': 150, 'sizes': [10, 5], 'text': FULL_D3055, 'price': 16.96},
    'PU 225': {'name': 'WAKOL PU 225 (klej)', 'usage': 1250, 'sizes': [10], 'text': "", 'price': 13.55},
    'MS 230': {'name': 'WAKOL MS 230 (klej)', 'usage': 1350, 'sizes': [18], 'text': "", 'price': 15.00},
    'MS 230 (B11 cement)': {'name': 'WAKOL MS 230 (klej)', 'usage': 1150, 'sizes': [18], 'text': "", 'price': 15.00},
    'MS 230 (B5 masa)': {'name': 'WAKOL MS 230 (klej)', 'usage': 900, 'sizes': [18], 'text': "", 'price': 15.00},
    'MS 260': {'name': 'WAKOL MS 260 (klej)', 'usage': 1350, 'sizes': [18], 'text': "", 'price': 13.46},
    'D 3318': {'name': '{PRODUCTS['D 3318']['name']} (klej)', 'usage': 350, 'sizes': [13], 'text': "", 'price': 17.60},
    'Z 645': {'name': 'WAKOL Z 645 (masa naprawcza)', 'usage': 1.6, 'sizes': [25], 'text': "", 'price': 4.26},
    'Z 645 (bruzdowane)': {'name': 'WAKOL Z 645 (masa szpachlowa)', 'usage': 2000, 'sizes': [25], 'text': "", 'price': 4.26},
    'Z 625': {'name': 'WAKOL Z 625 (masa samorozlewna)', 'usage_per_mm': 1.6, 'sizes': [25], 'text': FULL_Z625, 'price': 4.09},
    'Z 635': {'name': 'WAKOL Z 635 (masa samorozlewna)', 'usage_per_mm': 1.6, 'sizes': [25], 'text': FULL_Z635, 'price': 3.18},
    'Z 675': {'name': 'WAKOL Z 675 (masa samorozlewna)', 'usage_per_mm': 1.6, 'sizes': [25], 'text': FULL_Z675, 'price': 2.81},
    'D 3004 (bruzdowane)': {'name': 'WAKOL D 3004 (koncentrat)', 'usage': 75, 'sizes': [10, 5], 'text': "", 'price': 23.32},
    'AR 150': {'name': 'WAKOL AR 150 (mata kompensacyjna)', 'usage': 1000, 'sizes': [50], 'text': "", 'price': 8.06},
    'EM 140': {'name': 'WAKOL EM 140 (włóknina odsprzęgająca)', 'usage': 1000, 'sizes': [50], 'text': "", 'price': 17.67},
    'D 3060': {'name': 'WAKOL D 3060 (plastyfikator)', 'usage': 1000, 'sizes': [10], 'text': "", 'price': 26.00},
    'PU 280 (RP)': {'name': 'WAKOL PU 280 (grunt dla RP)', 'usage': 200, 'sizes': [11, 5], 'text': "* Zalecamy zagruntowanie całej powierzchni podłoża gruntówką wzmacniającą **{PRODUCTS['PU 280 (1W)']['name']}**. Aplikować wałkiem. Nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie ok. 200 g/m². Czas schnięcia 1 godzina. Czas do montażu – 72 godziny.", 'price': 54.27},
    'Płyta RP': {'name': 'WAKOL RP 704 (płyta odprzęgająca)', 'usage': 1000, 'sizes': [1], 'unit': 'szt', 'text': "* Na tak przygotowane podłoże zalecamy przyklejenie płyty odprzęgającej o grubości 4 mm **WAKOL RP 704**. Należy przyklejać klejem 2K PU (**{PRODUCTS['PU 225']['name']}**). Płytę odprzęgającą po ułożeniu należy docisnąć. Płytę można docinać używając noża trapezowego. Można układać parkiet, jeśli tylko klejona płyta nie przesuwa się w trakcie chodzenia po niej.", 'price': 22.05},
    'PS 205': {'name': 'WAKOL PS 205 (żywica lana)', 'sizes': [1], 'unit': 'kpl.', 'text': "", 'price': 48.40}
}

PRODUCTS_MAPEI = {
    k: dict(v, name=v['name'].replace('WAKOL', 'Mapei'), text='', price=0.0)
    for k, v in PRODUCTS_WAKOL.items()
}

_M_ECO_PU1K_1W = "* Zalecamy wykonanie gruntowania wzmacniającego poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **ECO PRIM PU 1K TURBO**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie **ok. 200 g/m²**. **Czas schnięcia – 2 godziny**."
_M_ECO_PU1K_BARRIER = "* Z uwagi na podwyższoną wilgotność zalecamy stworzenie **bariery przeciwwilgociowej** poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **ECO PRIM PU 1K TURBO**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **200 g/m²**. **Czas schnięcia – 2 godziny**. 2. warstwa **150 g/m²** – **czas schnięcia – 2 godziny**. **Czas do klejenia: 72 godziny od zagruntowania**."
_M_ECO_PU1K_BARRIER_PLYTA = "* Z uwagi na grubość płyty fundamentowej zalecamy stworzenie **bariery przeciwwilgociowej** poprzez zagruntowanie powierzchni podłoża gruntówką poliuretanową **ECO PRIM PU 1K TURBO**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **200 g/m²**. **Czas schnięcia – 2 godziny**. 2. warstwa **150 g/m²** – **czas schnięcia – 2 godziny**. **Czas do klejenia: 72 godziny od zagruntowania**."
_M_ECO_PU1K_RP = "* Zalecamy zagruntowanie całej powierzchni podłoża gruntówką poliuretanową **ECO PRIM PU 1K TURBO**. Aplikować wałkiem. Nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie ok. 200 g/m². Czas schnięcia 2 godziny. Czas do montażu – 72 godziny."
_M_PROSFAS = "* Zalecamy aplikację gruntówki wzmacniającej **PROSFAS** wałkiem lub pędzlem. Zużycie **ok. 1,5 kg/m²**. Aplikując **PROSFAS** należy zwrócić uwagę, aby dobrze wchłaniał się w podłoże i unikać powstawania kałuż na powierzchni jastrychu. **Po 7 dniach schnięcia** powierzchnię należy **przeszlifować papierem o gradacji 24 – 40** i dokładnie odkurzyć."
_M_ECO_PU1K_1W_S = "* Zalecamy wykonanie gruntowania wzmacniającego poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **ECO PRIM PU 1K**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie **ok. 200 g/m²**. **Czas schnięcia – 24 godziny**."
_M_ECO_PU1K_BARRIER_S = "* Zalecamy wykonanie **bariery przeciwwilgociowej** poprzez dwukrotne zagruntowanie gruntówką poliuretanową **ECO PRIM PU 1K**. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **200 g/m²**. **Czas schnięcia – 24 godziny**. 2. warstwa zużycie **150 g/m²**. **Czas schnięcia – 24 godziny**. **Czas klejenia 72 godziny od zagruntowania**."
_M_ECO_GRIP = "* Zalecamy zaaplikowanie mostka sczepnego **ECO PRIM GRIP PLUS**. Aplikować równomiernie wałkiem. Zużycie wynosi **ok. 200 g/m²**. **Czas schnięcia 1 godzina**."
_M_PRIMER_G_PRO = "* Zagruntować podłoże koncentratem gruntówki dyspersyjnej **MAPEI Primer G Pro**. Proporcje mieszania: 1 część koncentratu + 1 część wody. **Czas schnięcia 1 godzina**. Zużycie rozcieńczonej mikstury ok. 200 g/m², co daje zużycie koncentratu na poziomie **ok. 100 g/m²**."
_M_ECO_PRIM_T_PLUS = "* Zalecamy zagruntowanie całej powierzchni jastrychu gruntówką dyspersyjną **MAPEI Eco Prim T Plus** - aplikacja wałkiem **ok. 150 g/m²**. **Czas schnięcia ok. 1 godziny**."
_M_PLANOLIT_115 = "* Wylać masę wyrównawczą **MAPEI Planolit 115** - wymieszać ją w czystym naczyniu z zimną wodą w proporcji 5,75-6 litrów wody na 23 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,6 kg/m²/mm**. **Czas schnięcia 3mm - 24 godziny**."
_M_ULTRAPLAN_RENOVATION = "* Wylać masę wyrównawczą **MAPEI Ultraplan Renovation** – wymieszać ją w czystym naczyniu z zimną wodą w proporcji 4,5 litrów wody na 23 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,6 kg/m²/mm**. **Czas schnięcia 3mm - 24 godziny**."
_M_ULTRAPLAN_MAXI = "* Wylać masę wyrównawczą **MAPEI Ultraplan Maxi** – wymieszać ją w czystym naczyniu z zimną wodą w proporcji 4,75- 5,0 litrów wody na 23 kg masy. Mieszać unikając tworzenia się grudek. Prędkość obrotowa mieszadła może wynosić max. 600 obrotów na minutę. Wymieszaną masę nanosić w żądanej grubości na podłoże przy pomocy szpachli, łaty lub rakli. Przed pracą należy zwrócić uwagę na obecność wypełnień fug przy ścianach. Zużycie **ok. 1,7 kg/m²/mm**. **Czas schnięcia 3mm - 24 godziny**."

PRODUCTS_MAPEI = {
    k: dict(v, name=v['name'].replace('WAKOL', 'Mapei'), text='', price=0.0)
    for k, v in PRODUCTS_WAKOL.items()
}

PRODUCTS_MAPEI.update({
    'PU 280 (1W)':          {'name': 'MAPEI ECO PRIM PU 1K TURBO (1 warstwa)', 'usage': 200, 'sizes': [10], 'text': _M_ECO_PU1K_1W,           'price': 0.0},
    'PU 280 (Bariera)':     {'name': 'MAPEI ECO PRIM PU 1K TURBO (bariera)',    'usage': 350, 'sizes': [10], 'text': _M_ECO_PU1K_BARRIER,       'price': 0.0},
    'PU 280 (Bariera Płyta)':{'name': 'MAPEI ECO PRIM PU 1K TURBO (bariera)',   'usage': 350, 'sizes': [10], 'text': _M_ECO_PU1K_BARRIER_PLYTA, 'price': 0.0},
    'PU 280 (RP)':          {'name': 'MAPEI ECO PRIM PU 1K TURBO (grunt dla RP)','usage': 200, 'sizes': [10], 'text': _M_ECO_PU1K_RP,           'price': 0.0},
    'PU 235 (1W)':          {'name': 'MAPEI ECO PRIM PU 1K (1 warstwa)', 'usage': 200, 'sizes': [10, 5], 'text': _M_ECO_PU1K_1W_S,      'price': 0.0},
    'PU 235 (Bariera)':     {'name': 'MAPEI ECO PRIM PU 1K (bariera)',    'usage': 350, 'sizes': [10, 5], 'text': _M_ECO_PU1K_BARRIER_S, 'price': 0.0},
    'PS 275':               {'name': 'MAPEI PROSFAS',                      'usage': 1500, 'sizes': [25],  'text': _M_PROSFAS,            'price': 0.0},
    'D 3045':               {'name': 'MAPEI ECO PRIM GRIP PLUS',           'usage': 200,  'sizes': [10, 5, 1], 'text': _M_ECO_GRIP,     'price': 0.0},
    'D 3004':               {'name': 'MAPEI Primer G Pro',                 'usage': 100,  'sizes': [20, 10, 5, 1], 'unit': 'L', 'text': _M_PRIMER_G_PRO, 'price': 0.0},
    'D 3004 (bruzdowane)':  {'name': 'MAPEI Primer G Pro (koncentrat)', 'usage': 100,  'sizes': [20, 10, 5, 1], 'unit': 'L', 'text': "", 'price': 0.0},
    'D 3055':               {'name': 'MAPEI Eco Prim T Plus',              'usage': 150,  'sizes': [20, 5], 'text': _M_ECO_PRIM_T_PLUS, 'price': 0.0},
    'Z 675':                {'name': 'MAPEI Planolit 115 (masa samorozlewna)', 'usage_per_mm': 1.6, 'sizes': [23], 'text': _M_PLANOLIT_115, 'price': 0.0},
    'Z 635':                {'name': 'MAPEI Ultraplan Renovation (masa samorozlewna)', 'usage_per_mm': 1.6, 'sizes': [23], 'text': _M_ULTRAPLAN_RENOVATION, 'price': 0.0},
    'Z 625':                {'name': 'MAPEI Ultraplan Maxi (masa samorozlewna)', 'usage_per_mm': 1.7, 'sizes': [23], 'text': _M_ULTRAPLAN_MAXI, 'price': 0.0},
    'Z 645':                {'name': 'MAPEI Nivo Rapid (masa naprawcza)',  'usage': 1.6,  'sizes': [25], 'text': "", 'price': 0.0},
    'Z 645 (bruzdowane)':   {'name': 'MAPEI Nivo Rapid (masa szpachlowa)','usage': 2000, 'sizes': [25], 'text': "", 'price': 0.0},
    'D 3060':               {'name': 'MAPEI Latex Plus (plastyfikator)',   'usage': 1000, 'sizes': [10], 'unit': 'kg', 'text': "", 'price': 0.0},
    'PU 225':               {'name': 'MAPEI Ultrabond ECO P909 2K (klej)', 'usage': 1250, 'sizes': [10], 'text': "", 'price': 0.0},
    'MS 260':               {'name': 'MAPEI Ultrabond S965 1K (klej)',     'usage': 1350, 'sizes': [15], 'text': "", 'price': 0.0},
    'MS 230':               {'name': 'MAPEI Ultrabond ECO S948 1K (klej)', 'usage': 1350, 'sizes': [15], 'text': "", 'price': 0.0},
    'MS 230 (B11 cement)':  {'name': 'MAPEI Ultrabond ECO S948 1K (klej)', 'usage': 1150, 'sizes': [15], 'text': "", 'price': 0.0},
    'MS 230 (B5 masa)':     {'name': 'MAPEI Ultrabond ECO S948 1K (klej)', 'usage': 900,  'sizes': [15], 'text': "", 'price': 0.0},
    'PS 205':               {'name': 'MAPEI Epo Grip (żywica lana)',      'sizes': [10, 2], 'unit': 'kg', 'text': "", 'price': 0.0},
    'AR 150':               {'name': 'MAPEI MAPETHERM NET 150 (siatka zbrojeniowa)', 'usage': 1000, 'sizes': [50], 'text': "", 'price': 0.0},
})

_M_PROSFAS = "* Zalecamy aplikację gruntówki wzmacniającej **PROSFAS** wałkiem lub pędzlem. Zużycie **ok. 1,5 kg/m²**. Aplikując **PROSFAS** należy zwrócić uwagę, aby dobrze wchłaniał się w podłoże i unikać powstawania kałuż na powierzchni jastrychu. **Po 7 dniach schnięcia** powierzchnię należy **przeszlifować papierem o gradacji 24 – 40** i dokładnie odkurzyć."

_M_ECO_PU1K_1W_S = "* Zalecamy wykonanie gruntowania wzmacniającego poprzez zagruntowanie powierzchni jastrychu gruntówką poliuretanową **ECO PRIM PU 1K**. Aplikować wałkiem. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. Zużycie **ok. 200 g/m²**. **Czas schnięcia – 24 godziny**."
_M_ECO_PU1K_BARRIER_S = "* Zalecamy wykonanie **bariery przeciwwilgociowej** poprzez dwukrotne zagruntowanie gruntówką poliuretanową **ECO PRIM PU 1K**. Podczas aplikacji nie zostawiać kałuż tj. zbierać nadmiar niewchłoniętej gruntówki. 1. warstwa nałożona wałkiem **200 g/m²**. **Czas schnięcia – 24 godziny**. 2. warstwa zużycie **150 g/m²**. **Czas schnięcia – 24 godziny**. **Czas klejenia 72 godziny od zagruntowania**."

_M_ECO_GRIP = "* Zalecamy zaaplikowanie mostka sczepnego **ECO PRIM GRIP PLUS**. Aplikować równomiernie wałkiem. Zużycie wynosi **ok. 200 g/m²**. **Czas schnięcia 1 godzina**."

PRODUCTS_MAPEI.update({
    'PU 235 (1W)':      {'name': 'Mapei ECO PRIM PU 1K (1 warstwa)', 'usage': 200, 'sizes': [10, 5], 'text': _M_ECO_PU1K_1W_S,      'price': 0.0},
    'PU 235 (Bariera)': {'name': 'Mapei ECO PRIM PU 1K (bariera)',    'usage': 350, 'sizes': [10, 5], 'text': _M_ECO_PU1K_BARRIER_S, 'price': 0.0},
    'PS 275':           {'name': 'Mapei PROSFAS',                      'usage': 1500, 'sizes': [25],  'text': _M_PROSFAS,            'price': 0.0},
    'D 3045':           {'name': 'Mapei ECO PRIM GRIP PLUS',           'usage': 200,  'sizes': [10, 5, 1], 'text': _M_ECO_GRIP,     'price': 0.0},
})
PRODUCTS_UZIN = {
    k: dict(v, name=v['name'].replace('WAKOL', 'Uzin'), text='', price=0.0)
    for k, v in PRODUCTS_WAKOL.items()
}

PRODUCTS = PRODUCTS_WAKOL

def _calc_combo(needed_kg, sizes, unit):
    if not sizes:
        q = math.ceil(needed_kg)
        return q, f"{q} {unit}"
    if len(sizes) == 1:
        q = math.ceil(needed_kg / sizes[0])
        return q * sizes[0], f"{q}x {sizes[0]} {unit}"
    large, small = sizes[0], sizes[1]
    best_waste = None
    best_nl, best_ns = 0, 1
    for nl in range(0, math.ceil(needed_kg / large) + 1):
        rem = needed_kg - nl * large
        ns = math.ceil(rem / small) if rem > 0 else 0
        total = nl * large + ns * small
        waste = total - needed_kg
        if waste >= -0.001:
            if best_waste is None or waste < best_waste - 0.001 or (abs(waste - best_waste) < 0.001 and nl + ns < best_nl + best_ns):
                best_waste, best_nl, best_ns = waste, nl, ns
    if best_nl == 0 and best_ns == 0:
        best_ns = 1
    bought = best_nl * large + best_ns * small
    parts = []
    if best_nl > 0: parts.append(f"{best_nl}x {large} {unit}")
    if best_ns > 0: parts.append(f"{best_ns}x {small} {unit}")
    if not parts:
        bought = small
        parts = [f"1x {small} {unit}"]
    return bought, " + ".join(parts)

def write_and_track(dane, rep, prod_key, custom_kg=None):
    prod = PRODUCTS[prod_key]

    if 'written_texts' not in dane:
        dane['written_texts'] = set()

    if prod_key not in dane['written_texts']:
        if prod['text']:
            rep.write(prod['text'])
        dane['written_texts'].add(prod_key)

    if 'materials' not in dane:
        dane['materials'] = []

    needed_kg = 0
    if custom_kg is not None:
        needed_kg = custom_kg
    else:
        area = dane.get('area_m2')
        if not area: return
        if 'usage_per_mm' in prod:
            thick = dane.get('leveling_thickness')
            if not thick: return
            needed_kg = area * thick * prod['usage_per_mm']
        else:
            needed_kg = (area * prod['usage']) / 1000.0

    if needed_kg <= 0: return
    sizes = sorted(prod['sizes'], reverse=True)
    unit = prod.get('unit', 'kg')
    price = prod.get('price', 0)
    pkg_size = sizes[0] if sizes else 1

    existing = next((m for m in dane['materials'] if m['name'] == prod['name']), None)
    if existing:
        new_total = round(existing['kg'] + needed_kg, 2)
        new_bought, new_combo = _calc_combo(new_total, sizes, unit)
        existing['kg'] = new_total
        existing['bought_qty'] = new_bought
        existing['combo'] = new_combo
        existing['total_cost'] = new_bought * price
        existing['exact_cost'] = new_total * price
    else:
        bought_qty, combo = _calc_combo(needed_kg, sizes, unit)
        dane['materials'].append({
            'name': prod['name'],
            'kg': round(needed_kg, 2),
            'bought_qty': bought_qty,
            'combo': combo,
            'unit': unit,
            'pkg_size': pkg_size,
            'price_per_unit': price,
            'total_cost': bought_qty * price,
            'exact_cost': needed_kg * price,
        })

def _fmt_pkg(needed, pkg_size, unit):
    real_p = needed / pkg_size if pkg_size else needed
    buy_p = math.ceil(real_p)
    real_str = str(int(real_p)) if real_p == int(real_p) else f"{real_p:.1f}".replace('.', ',')
    if buy_p == 1:
        buy_word = "opakowanie"
    elif buy_p in [2, 3, 4]:
        buy_word = "opakowania"
    else:
        buy_word = "opakowań"
    return f"{real_str} ({buy_p} {buy_word} po {pkg_size} {unit})"

def _kosztorys_line_v1(m):
    unit = m.get('unit', 'kg')
    pkg_size = m.get('pkg_size', 1)
    price = m.get('price_per_unit', 0)
    kg = m['kg']
    cost = kg * price
    real_p = kg / pkg_size if pkg_size else kg
    real_str = str(int(real_p)) if real_p == int(real_p) else f"{real_p:.1f}".replace('.', ',')
    return f"- {m['name']}: {kg} {unit}({real_str} opak.{pkg_size}{unit}) x {price:.2f} PLN = **{cost:.2f} PLN**", cost

def _kosztorys_line_v2(m):
    unit = m.get('unit', 'kg')
    price = m.get('price_per_unit', 0)
    bought = m['bought_qty']
    cost = bought * price
    combo = m.get('combo', f"{bought} {unit}")
    return f"- {m['name']}: {bought} {unit}({combo}) x {price:.2f} PLN = **{cost:.2f} PLN**", cost

def render_potrzebne_materialy(dane, rep):
    if not dane.get('area_m2'): return
    if not dane.get('materials'): return
    if not dane.get('include_cost', False): return

    rep.write("\n**Wariant 1: Kosztorys materiałowy – rzeczywiste zużycie (Netto)**")
    total1 = 0.0
    for m in dane['materials']:
        if m.get('price_per_unit', 0) <= 0: continue
        line, cost = _kosztorys_line_v1(m)
        rep.write(line)
        total1 += cost
    rep.write(f"**RAZEM NETTO (Wariant 1): {total1:.2f} PLN**")

    rep.write("\n**Wariant 2: Kosztorys materiałowy – pełne opakowania (Netto)**")
    total2 = 0.0
    for m in dane['materials']:
        if m.get('price_per_unit', 0) <= 0: continue
        line, cost = _kosztorys_line_v2(m)
        rep.write(line)
        total2 += cost
    rep.write(f"**RAZEM NETTO (Wariant 2): {total2:.2f} PLN**")

def render_wspolne_zalecenia_podloze(dane, rep):
    rep.write("**a) przygotowanie podłoża:**")
    if dane.get('requires_demolition'):
        rep.write("* **Demontaż starej okładziny.**")
    if dane['dilatations_obw_ok'] == "NIE":
        rep.write("* Odtworzenie dylatacji obwodowych.")
    if dane['cracks_klaw'] == "TAK":
        rep.write("* Rozbruzdowanie klawiszujących dylatacji pozornych.")
    if dane['cracks_pek'] == "TAK":
        rep.write("* Rozbruzdowanie pęknięć wymagających zespolenia.")
    
    if dane['substrate'] == "płytki ceramiczne":
        rep.write("* Mechaniczne usunięcie szkliwa płytek poprzez szlif.")
        rep.write("* Sprawdzenie stabilności połączenia płytek z podłożem (głuche elementy skuć i zaszpachlować masą szpachlową).")
    elif dane['substrate'] == "podłoże drewniane (parkiet, deska)":
        rep.write("* **Szlif podłoża** w celu wyrównania i oczyszczenia powierzchni drewnianej.")
    elif dane['substrate'] == "podłoże z płyty OSB":
        rep.write("* **Szlif podłoża** w celu wyrównania i oczyszczenia powierzchni płyty OSB.")
        rep.write("* Sprawdzenie i dokręcenie wkrętów mocujących płyty OSB (łby wkrętów muszą być zagłębione w powierzchni).")
    else:
        rep.write("* **Szlif podłoża** w celu uzyskania porowatej i chłonnej powierzchni!")
        
    rep.write("* Dokładne odkurzenie powierzchni odkurzaczem przemysłowym.")
    
    if dane['curing_not_done']:
        if dane['is_moisture_neg']:
            if dane['decision_after_cure'] == "osuszanie przed barierą":
                rep.write(f"* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** w celu obniżenia poziomu wilgoci do max. **{dane.get('barrier_max', '2.8')}%**, a następnie wykonanie bariery.")
            else:
                rep.write(f"* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** w celu uzyskania normatywnego poziomu wilgoci **{dane['norm_val_bracket']}**.")
        else:
            rep.write(f"* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** zgodnie z protokołem.")
    elif dane['is_moisture_neg']:
        if dane['decision_after_cure'] == "Wykonanie bariery przeciwwilgociowej":
            rep.write("* Zalecamy wykonanie **bariery przeciwwilgociowej**.")
        elif dane['decision_after_cure'] == "osuszanie przed barierą":
            rep.write(f"* Zalecamy doprowadzenie poziomu wilgoci do max. **{dane.get('barrier_max', '2.8')}%** poprzez dalsze osuszanie, a następnie wykonanie bariery przeciwwilgociowej.")
        else:
            rep.write(f"* Zalecamy doprowadzenie do normatywnego poziomu wilgoci **{dane['norm_val_bracket']}** poprzez {dane['decision_after_cure']}.")

    rep.write("**b) naprawa i wzmocnienie podłoża:**")
    if dane['curing_not_done']:
        if dane['is_moisture_neg']:
            if dane['decision_after_cure'] == "osuszanie przed barierą":
                rep.write(f"Po doprowadzeniu poziomu wilgoci do max. **{dane.get('barrier_max', '2.8')}%** poprzez **przeprowadzenie procesu wygrzewania** zalecamy:")
            else:
                rep.write(f"Po doprowadzeniu do normatywnego poziomu wilgoci **{dane['norm_val_bracket']}** jastrychu poprzez **przeprowadzenie procesu wygrzewania** zalecamy:")
        else:
            rep.write("Po **przeprowadzeniu pełnego procesu wygrzewania** zalecamy:")
    elif dane['needs_drying_action']:
        if dane['decision_after_cure'] == "osuszanie przed barierą":
            rep.write(f"Po doprowadzeniu poziomu wilgoci do max. **{dane.get('barrier_max', '2.8')}%** zalecamy:")
        elif dane['decision_after_cure'] == "kolejny proces wygrzewania":
            rep.write(f"Po doprowadzeniu do normatywnego poziomu wilgoci **{dane['norm_val_bracket']}** poprzez przeprowadzenie kolejnego procesu wygrzewania zalecamy:")
        else:
            rep.write(f"Po doprowadzeniu do normatywnego poziomu wilgoci **{dane['norm_val_bracket']}** zalecamy:")
    
    if dane['cracks_klaw'] == "TAK" or dane['cracks_pek'] == "TAK":
        if dane['strength_val'] == 1 and dane['substrate'] != "jastrych anhydrytowy":
            write_and_track(dane, rep, 'PS 275')
        
        firma_is_mapei = (dane.get('firma') == "Mapei")
        if firma_is_mapei:
            rep.write(f"* Pęknięcia / Klawiszujące dylatacje - zespolić żywicą laną **{PRODUCTS['PS 205']['name']}**. Wymieszaną żywicę wlewać w pęknięcia, nadmiar zgarnąć lub zatrzeć.")
        else:
            rep.write(f"* Pęknięcia / Klawiszujące dylatacje - zespolić żywicą laną **{PRODUCTS['PS 205']['name']}**. Wymieszaną żywicę wlewać w pęknięcia, nadmiar zgarnąć lub zatrzeć.")
            
        total_meters = 0
        if dane['cracks_klaw'] == "TAK":
            total_meters += dane.get('klaw_meters') or 0
        if dane['cracks_pek'] == "TAK":
            total_meters += dane.get('pek_meters') or 0
        if total_meters > 0:
            if firma_is_mapei:
                write_and_track(dane, rep, 'PS 205', custom_kg=total_meters / 7.5)
            else:
                write_and_track(dane, rep, 'PS 205', custom_kg=total_meters / 6.5)

    if dane['holes'] == "TAK":
        kg_z645 = None
        if dane.get('holes_width') and dane.get('holes_length') and dane.get('holes_depth'):
            area_h = (dane['holes_width'] / 100.0) * dane['holes_length']
            thick_mm = dane['holes_depth'] * 10.0
            kg_z645 = area_h * thick_mm * 1.6

        firma_is_mapei = (dane.get('firma') == "Mapei")
        if dane.get('holes_depth') and dane['holes_depth'] >= 1.0:
            if kg_z645 is not None: kg_z645 /= 2.0
            if firma_is_mapei:
                rep.write(f"* Ubytki zaszpachlować masą **{PRODUCTS['Z 645']['name']}** wymieszaną z piaskiem kwarcowym (dodatek 30%). Czas schnięcia przed klejeniem 12 godzin.")
            else:
                rep.write(f"* Ubytki zaszpachlować masą **{PRODUCTS['Z 645']['name']}** wymieszaną z piaskiem kwarcowym w proporcji 1:1  – czas schnięcia 1 godzina.")
        else:
            if firma_is_mapei:
                rep.write(f"* Ubytki zaszpachlować masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** (8 kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Ubytki zaszpachlować masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** (7 litrów na 25 kg masy). Czas schnięcia min. 3h. W razie potrzeby użyć siatki zbrojeniowej {PRODUCTS['AR 150']['name']}.")
        
        if kg_z645 is not None:
            write_and_track(dane, rep, 'Z 645', custom_kg=kg_z645)

    if dane.get('local_leveling') == "TAK" and dane.get('local_leveling_kg'):
        details = dane.get('local_leveling_details', '')
        firma_is_mapei = (dane.get('firma') == "Mapei")
        ratio = 8.0 if firma_is_mapei else 7.0
        
        if firma_is_mapei:
            rep.write(f"* Miejscowe wyrównanie podłoża{details} masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
        else:
            rep.write(f"* Miejscowe wyrównanie podłoża{details} masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
            
        write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=dane['local_leveling_kg'])
        bags_local = dane['local_leveling_kg'] / 25.0
        write_and_track(dane, rep, 'D 3060', custom_kg=bags_local * ratio)

    if dane['heating_exists'] == "TAK" and dane['h_type'] == "bruzdowane":
        if dane['bruzdowane_wybor'] == "masa samorozlewna":
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Podłoże zagruntować koncentratem gruntówki dyspersyjnej **{PRODUCTS['D 3004']['name']}**. Proporcje mieszania: 1 część koncentratu + 1 część wody; Czas schnięcia: 1h. Sposób nanoszenia: wałek do gruntowania. Zużycie: ok. 100 g/m² koncentratu.")
            else:
                rep.write(f"* Podłoże zagruntować koncentratem gruntówki dyspersyjnej **{PRODUCTS['D 3004']['name']}**. Proporcje mieszania: 1 część koncentratu + 1 część wody; Czas schnięcia: 1h. Sposób nanoszenia: wałek do gruntowania microfazer. Zużycie: ok. 75 g/m² koncentratu.")
            write_and_track(dane, rep, 'D 3004 (bruzdowane)')
            
            if firma_is_mapei:
                rep.write(f"* Na tak przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na tak przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
            
            area = dane.get('area_m2') or 0
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_bruzdowane = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_bruzdowane)
                bags_z645 = math.ceil(kg_z645_bruzdowane / 25.0)
                write_and_track(dane, rep, 'D 3060', custom_kg=bags_z645 * ratio)
            write_and_track(dane, rep, 'Z 635', custom_kg=area * 5 * 1.5)
        elif dane['bruzdowane_wybor'] == "płyta RP":
            area = dane.get('area_m2') or 0
            write_and_track(dane, rep, 'PU 280 (RP)')
            write_and_track(dane, rep, 'Płyta RP', custom_kg=math.ceil(area / 0.6) if area > 0 else 0)
            write_and_track(dane, rep, 'PU 225')

def render_wspolna_chemia(dane, rep):
    used_d3004 = False
    if dane.get('h_type') == "bruzdowane" and dane.get('bruzdowane_wybor'):
        return True # Pomijamy standardową chemię, obsłużona w naprawie podłoża

    if dane['substrate'] in ["płytki ceramiczne", "podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB"]:
        return False

    if dane['decision_after_cure'] in ["Wykonanie bariery przeciwwilgociowej", "osuszanie przed barierą"]:
        if dane['substrate'] == "płyta fundamentowa":
            write_and_track(dane, rep, 'PU 280 (Bariera Płyta)')
        elif dane['strength_val'] <= 2: write_and_track(dane, rep, 'PU 235 (Bariera)')
        else: write_and_track(dane, rep, 'PU 280 (Bariera)')
    elif not dane['decision_after_cure'] or "Wykonanie" not in str(dane['decision_after_cure']):
        if dane.get('has_adhesive_residues'):
            write_and_track(dane, rep, 'PU 280 (1W)')
            return False
        if dane['needs_levelling'] == "TAK":
            if dane['strength_val'] in [3, 4, 5]:
                if dane['substrate'] == "jastrych anhydrytowy" and dane['leveling_thickness'] and dane['leveling_thickness'] > 5:
                    write_and_track(dane, rep, 'PU 280 (1W)')
                else:
                    write_and_track(dane, rep, 'D 3004')
                    used_d3004 = True
            else:
                if dane['strength_val'] == 1:
                    if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                    else:
                        write_and_track(dane, rep, 'PS 275')
                        if dane.get('firma') != 'Mapei':
                            write_and_track(dane, rep, 'PU 280 (1W)')
                elif dane['strength_val'] == 2: write_and_track(dane, rep, 'PU 280 (1W)')
        else:
            if dane['strength_val'] == 1:
                if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PS 275')
            elif dane['strength_val'] == 2: write_and_track(dane, rep, 'PU 235 (1W)')
            elif dane['strength_val'] in [3, 4]: write_and_track(dane, rep, 'PU 280 (1W)')
    return used_d3004

def render_chemia_deska_warstwowa(dane, rep):
    used_d3004 = False
    if dane.get('h_type') == "bruzdowane" and dane.get('bruzdowane_wybor'):
        return True

    if dane['substrate'] in ["płytki ceramiczne", "podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB"]:
        return False

    if dane['decision_after_cure'] in ["Wykonanie bariery przeciwwilgociowej", "osuszanie przed barierą"]:
        if dane['substrate'] == "płyta fundamentowa":
            write_and_track(dane, rep, 'PU 280 (Bariera Płyta)')
        elif dane['strength_val'] <= 2: write_and_track(dane, rep, 'PU 235 (Bariera)')
        else: write_and_track(dane, rep, 'PU 280 (Bariera)')
    elif not dane['decision_after_cure'] or "Wykonanie" not in str(dane['decision_after_cure']) and "barierą" not in str(dane['decision_after_cure']):
        if dane.get('has_adhesive_residues'):
            write_and_track(dane, rep, 'PU 280 (1W)')
            return False
        if dane['needs_levelling'] == "TAK":
            if dane['strength_val'] in [3, 4, 5]:
                if dane['substrate'] == "jastrych anhydrytowy" and dane['leveling_thickness'] and dane['leveling_thickness'] > 5:
                    write_and_track(dane, rep, 'PU 280 (1W)')
                else:
                    write_and_track(dane, rep, 'D 3004')
                    used_d3004 = True
            elif dane['strength_val'] == 2:
                write_and_track(dane, rep, 'PU 280 (1W)')
            elif dane['strength_val'] == 1:
                if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PS 275')
                    write_and_track(dane, rep, 'PU 280 (1W)')
        else:
            if dane['strength_val'] == 1:
                if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PS 275')
            elif dane['strength_val'] == 2: write_and_track(dane, rep, 'PU 235 (1W)')
            elif dane['strength_val'] == 3: write_and_track(dane, rep, 'PU 280 (1W)')
            elif dane['strength_val'] in [4, 5]: write_and_track(dane, rep, 'D 3055')
    return used_d3004

def render_chemia_deska_lita(dane, rep):
    used_d3004 = False
    if dane.get('h_type') == "bruzdowane" and dane.get('bruzdowane_wybor'):
        return True

    if dane['substrate'] in ["płytki ceramiczne", "podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB"]:
        return False

    if dane['decision_after_cure'] in ["Wykonanie bariery przeciwwilgociowej", "osuszanie przed barierą"]:
        if dane['substrate'] == "płyta fundamentowa":
            write_and_track(dane, rep, 'PU 280 (Bariera Płyta)')
        elif dane['strength_val'] <= 2: write_and_track(dane, rep, 'PU 235 (Bariera)')
        else: write_and_track(dane, rep, 'PU 280 (Bariera)')
    elif not dane['decision_after_cure'] or "Wykonanie" not in str(dane['decision_after_cure']) and "barierą" not in str(dane['decision_after_cure']):
        if dane.get('has_adhesive_residues'):
            write_and_track(dane, rep, 'PU 280 (1W)')
            return False
        if dane['needs_levelling'] == "TAK":
            if dane['strength_val'] in [3, 4, 5]:
                if dane['substrate'] == "jastrych anhydrytowy" and dane['leveling_thickness'] and dane['leveling_thickness'] > 5:
                    write_and_track(dane, rep, 'PU 280 (1W)')
                else:
                    write_and_track(dane, rep, 'D 3004')
                    used_d3004 = True
            else:
                if dane['strength_val'] == 1:
                    if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                    else:
                        write_and_track(dane, rep, 'PS 275')
                        if dane.get('firma') != 'Mapei':
                            write_and_track(dane, rep, 'PU 280 (1W)')
                elif dane['strength_val'] == 2: write_and_track(dane, rep, 'PU 280 (1W)')
        else:
            if dane['strength_val'] == 1:
                if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PS 275')
            elif dane['strength_val'] == 2: write_and_track(dane, rep, 'PU 235 (1W)')
            elif dane['strength_val'] in [3, 4]: write_and_track(dane, rep, 'PU 280 (1W)')
            elif dane['strength_val'] == 5: write_and_track(dane, rep, 'D 3055')
    return used_d3004

# --- SEKCJA: DESKA WARSTWOWA ---
def generate_report_deska_warstwowa(dane, rep):
    render_wspolne_dane_optyczne(dane, rep)
    
    if dane['flooring_type'] == "deska warstwowa":
        nazwa_okladziny, tytul_sekcji = "podłogę drewnianą", "Deska Warstwowa"
    elif dane['flooring_type'] == "lity parkiet (maks. 8 cm x 60 cm)":
        nazwa_okladziny, tytul_sekcji = "lity parkiet", "Lity Parkiet (maks. 8 cm x 60 cm)"
    elif dane['flooring_type'] == "mozaika drewniana (min. 16 mm grubości, maks. 20 cm długości)":
        nazwa_okladziny, tytul_sekcji = "mozaikę drewnianą", "Mozaika Drewniana (min. 16 mm, maks. 20 cm)"
    else:
        nazwa_okladziny, tytul_sekcji = "podłogę laminowaną", "Podłoga laminowana"
    
    if dane['substrate'] == "jastrych cementowy":
        rep.write(f"**Aby bezpiecznie kleić {nazwa_okladziny} na jastrychu cementowym, jego wytrzymałość na ścinanie musi wynosić między 1,5 a 2,0 N/mm² a wilgotność nie może przekraczać 1,8% CM. (z ogrzewaniem podłogowym max. 1,5% CM).**")
    elif dane['substrate'] == "jastrych anhydrytowy":
        rep.write(f"**Aby bezpiecznie kleić {nazwa_okladziny} na jastrychu anhydrytowym zgodnie z wytycznymi ITB, jego wytrzymałość na ścinanie musi wynosić 2,0 N/mm² a wilgotność nie może przekraczać 0,5% CM. (z ogrzewaniem podłogowym max. 0,3% CM).**")
    
    rep.markdown(f"#### **II. Zalecenia techniczne ({tytul_sekcji})**")
    
    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_chemia_deska_warstwowa(dane, rep)

    if dane['needs_levelling'] == "TAK" and dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 635')

    rep.write("**c) klejenie okładziny:**")
    if dane['substrate'] == "płytki ceramiczne" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie {nazwa_okladziny} należy przeprowadzić przy użyciu dwuskładnikowego kleju poliuretanowego **{PRODUCTS['PU 225']['name']}** (szpachla B11, zużycie: 1250 g/m²).")
        write_and_track(dane, rep, 'PU 225')
    elif dane['substrate'] == "płyta fundamentowa" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie {nazwa_okladziny} należy przeprowadzić przy użyciu kleju polimerowego twardo-elastycznego **{PRODUCTS['MS 260']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 260')
    elif dane['substrate'] == "jastrych anhydrytowy" and dane['strength_val'] == 1:
        rep.write(f"Klejenie {nazwa_okladziny} należy przeprowadzić przy użyciu kleju do parkietu **{PRODUCTS['MS 230']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 230')
    elif dane.get('klej_typ') == "bezprzesuwny":
        rep.write(f"Klejenie {nazwa_okladziny} należy przeprowadzić przy użyciu kleju do parkietu **{PRODUCTS['PU 225']['name']}** (szpachla B11, zużycie: 1250 g/m²).")
        write_and_track(dane, rep, 'PU 225')
    else:
        rep.write(f"Klejenie {nazwa_okladziny} należy przeprowadzić przy użyciu kleju do parkietu **{PRODUCTS['MS 230']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 230')
        
    render_potrzebne_materialy(dane, rep)

# --- SEKCJA: DESKA LITA ---
def generate_report_deska_lita(dane, rep):
    render_wspolne_dane_optyczne(dane, rep)
    if dane['substrate'] == "jastrych cementowy":
        rep.write("**Aby bezpiecznie kleić podłogę drewnianą na jastrychu cementowym, jego wytrzymałość na ścinanie musi wynosić między 1,5 a 2,0 N/mm² a wilgotność nie może przekraczać 1,8% CM. (z ogrzewaniem podłogowym max. 1,5% CM).**")
    elif dane['substrate'] == "jastrych anhydrytowy":
        rep.write("**Aby bezpiecznie kleić podłogę drewnianą na jastrychu anhydrytowym zgodnie z wytycznymi ITB, jego wytrzymałość na ścinanie musi wynosić 2,0 N/mm² a wilgotność nie może przekraczać 0,5% CM. (z ogrzewaniem podłogowym max. 0,3% CM).**")
    rep.markdown("#### **II. Zalecenia techniczne (Deska Lita)**")
    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_chemia_deska_lita(dane, rep)

    if dane['needs_levelling'] == "TAK" and dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 625')

    rep.write("**c) klejenie okładziny:**")
    if dane['substrate'] == "płytki ceramiczne" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie podłogi z deski litej należy przeprowadzić przy użyciu dwuskładnikowego kleju poliuretanowego **{PRODUCTS['PU 225']['name']}** (szpachla B11, zużycie: 1250 g/m²).")
        write_and_track(dane, rep, 'PU 225')
    elif dane['substrate'] == "płyta fundamentowa" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie podłogi z deski litej należy przeprowadzić przy użyciu kleju polimerowego twardo-elastycznego **{PRODUCTS['MS 260']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 260')
    elif dane.get('klej_typ') == "bezprzesuwny":
        rep.write(f"Klejenie podłogi z deski litej należy przeprowadzić przy użyciu kleju **{PRODUCTS['PU 225']['name']}** (szpachla B11, zużycie: 1250 g/m²).")
        write_and_track(dane, rep, 'PU 225')
    else:
        rep.write(f"Klejenie podłogi z deski litej należy przeprowadzić przy użyciu kleju polimerowego twardo-elastycznego **{PRODUCTS['MS 260']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 260')
    render_potrzebne_materialy(dane, rep)

# --- SEKCJA: LVT CIENKIE ---
def generate_report_lvt_cienkie(dane, rep):
    render_wspolne_dane_optyczne(dane, rep)
    rep.markdown("#### **II. Zalecenia techniczne (LVT Cienkie)**")
    
    if dane.get('already_levelled') == "TAK":
        rep.write("**a) przygotowanie podłoża:**")
        if dane['curing_not_done']:
            rep.write("* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** zgodnie z protokołem.")
            rep.write("Po **przeprowadzeniu pełnego procesu wygrzewania** zalecamy:")
        rep.write("* Szlif podłoża w celu uzyskania gładkiej powierzchni.")
        rep.write("* Dokładne odkurzenie powierzchni odkurzaczem przemysłowym.")
        rep.write("**b) klejenie okładziny:**")
        rep.write(f"Klejenie podłogi winylowej (LVT) należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3318']['name']} (szpachla TKB A2, zużycie: 350 g/m²). · Czas wstępnego odparowania: ok. 5 - 10 minut. · Czas układania: ok. 10 minut")
        write_and_track(dane, rep, 'D 3318')
        render_potrzebne_materialy(dane, rep)
        return

    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_wspolna_chemia(dane, rep)

    if dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 675')

    rep.write("* Po wyschnięciu masy samorozlewnej zalecamy szlif podłoża w celu uzyskania gładkiej powierzchni oraz dokładne odkurzenie.")

    rep.write("**c) klejenie okładziny:**")
    rep.write(f"Klejenie podłogi winylowej (LVT) należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3318']['name']} (szpachla TKB A2, zużycie: 350 g/m²). · Czas wstępnego odparowania: ok. 5 - 10 minut. · Czas układania: ok. 10 minut")
    write_and_track(dane, rep, 'D 3318')
    render_potrzebne_materialy(dane, rep)

# --- SEKCJA: LVT GRUBE ---
def render_chemia_lvt_grube(dane, rep):
    used_d3004 = False
    if dane.get('h_type') == "bruzdowane" and dane.get('bruzdowane_wybor'):
        return True

    if dane['substrate'] in ["płytki ceramiczne", "podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB"]:
        return False

    if dane['decision_after_cure'] in ["Wykonanie bariery przeciwwilgociowej", "osuszanie przed barierą"]:
        if dane['substrate'] == "płyta fundamentowa":
            write_and_track(dane, rep, 'PU 280 (Bariera Płyta)')
        elif dane['strength_val'] <= 2: write_and_track(dane, rep, 'PU 235 (Bariera)')
        else: write_and_track(dane, rep, 'PU 280 (Bariera)')
    elif not dane['decision_after_cure'] or ("Wykonanie" not in str(dane['decision_after_cure']) and "barierą" not in str(dane['decision_after_cure'])):
        if dane.get('has_adhesive_residues'):
            write_and_track(dane, rep, 'PU 280 (1W)')
            return False
        if dane['needs_levelling'] == "TAK":
            if dane['strength_val'] in [3, 4, 5]:
                if dane['substrate'] == "jastrych anhydrytowy" and dane['leveling_thickness'] and dane['leveling_thickness'] > 5:
                    write_and_track(dane, rep, 'PU 280 (1W)')
                else:
                    write_and_track(dane, rep, 'D 3004')
                    used_d3004 = True
            elif dane['strength_val'] == 2:
                write_and_track(dane, rep, 'PU 280 (1W)')
            elif dane['strength_val'] == 1:
                if dane['substrate'] == "jastrych anhydrytowy": write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PS 275')
                    write_and_track(dane, rep, 'PU 280 (1W)')
        else:
            if dane['strength_val'] in [3, 4, 5]:
                write_and_track(dane, rep, 'D 3055')
            elif dane['strength_val'] in [1, 2]:
                if dane['substrate'] == "jastrych anhydrytowy" and dane['strength_val'] == 1:
                    write_and_track(dane, rep, 'PU 235 (1W)')
                else:
                    write_and_track(dane, rep, 'PU 280 (1W)')
    return used_d3004

def generate_report_lvt_grube(dane, rep):
    render_wspolne_dane_optyczne(dane, rep)
    rep.markdown("#### **II. Zalecenia techniczne (LVT Grube z twardym rdzeniem)**")
    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_chemia_lvt_grube(dane, rep)

    if dane['needs_levelling'] == "TAK" and dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 675')

    rep.write("**c) klejenie okładziny:**")
    bottom_type = dane.get('lvt_bottom_type', '')
    if bottom_type == "Winyl na piankowym spodzie":
        rep.write("**Brak możliwości klejenia.** Możliwość montażu okładziny jedynie na pływająco.")
    elif dane['substrate'] == "płytki ceramiczne" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie podłogi LVT ({bottom_type}) należy przeprowadzić przy użyciu dwuskładnikowego kleju poliuretanowego **{PRODUCTS['PU 225']['name']}** (szpachla B11, zużycie: 1250 g/m²).")
        write_and_track(dane, rep, 'PU 225')
    elif dane['substrate'] == "płyta fundamentowa" and dane['needs_levelling'] == "NIE":
        rep.write(f"Klejenie podłogi LVT ({bottom_type}) należy przeprowadzić przy użyciu kleju polimerowego twardo-elastycznego **{PRODUCTS['MS 260']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 260')
    elif dane['needs_levelling'] == "TAK" or dane['substrate'] == "masa samorozlewna":
        rep.write(f"Klejenie podłogi LVT ({bottom_type}) na masie samorozlewnej należy przeprowadzić przy użyciu kleju **{PRODUCTS['MS 230']['name']}** (szpachla B5, zużycie: 900 g/m²).")
        write_and_track(dane, rep, 'MS 230 (B5 masa)')
    elif dane['substrate'] == "jastrych cementowy":
        rep.write(f"Klejenie podłogi LVT ({bottom_type}) na jastrychu cementowym należy przeprowadzić przy użyciu kleju **{PRODUCTS['MS 230']['name']}** (szpachla B11, zużycie: 1150 g/m²).")
        write_and_track(dane, rep, 'MS 230 (B11 cement)')
    else:
        rep.write(f"Klejenie podłogi LVT ({bottom_type}) należy przeprowadzić przy użyciu kleju **{PRODUCTS['MS 230']['name']}** (szpachla B13, zużycie: 1350 g/m²).")
        write_and_track(dane, rep, 'MS 230')
    render_potrzebne_materialy(dane, rep)

# --- SEKCJA: PCV W ROLCE ---
def generate_report_pcv_w_rolce(dane, rep):
    if dane['needs_levelling'] == "NIE" and dane['already_levelled'] == "NIE":
        rep.error("BŁĄD: Pod okładzinę PCV w rolce wymagane jest wyrównanie podłoża. Poinformuj klienta o konieczności wylania masy!")
        return
        
    render_wspolne_dane_optyczne(dane, rep)
    rep.markdown("#### **II. Zalecenia techniczne (PCV w rolce)**")
    
    if dane['already_levelled'] == "TAK":
        rep.write("**a) przygotowanie podłoża:**")
        if dane['curing_not_done']:
            rep.write("* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** zgodnie z protokołem.")
            rep.write("Po **przeprowadzeniu pełnego procesu wygrzewania** zalecamy:")
        rep.write("* Szlif podłoża w celu uzyskania gładkiej powierzchni.")
        rep.write("* Dokładne odkurzenie powierzchni odkurzaczem przemysłowym.")
        rep.write("**b) klejenie okładziny PCV:**")
        rep.write(f"Klejenie wykładziny PCV w rolce należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3307']['name']} (szpachla TKB A2, zużycie: 300 – 330 g/m²). · Czas wstępnego odparowania: ok. 10 - 20 minut. · Czas układania: ok. 15 - 20 minut")
        render_potrzebne_materialy(dane, rep)
        return

    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_wspolna_chemia(dane, rep)

    if dane['needs_levelling'] == "TAK" and dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 675')

    rep.write("* Po wyschnięciu masy samorozlewnej zalecamy szlif podłoża w celu uzyskania gładkiej powierzchni oraz dokładne odkurzenie.")

    rep.write("**c) klejenie okładziny PCV:**")
    rep.write(f"Klejenie wykładziny PCV w rolce należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3307']['name']} (szpachla TKB A2, zużycie: 300 – 330 g/m²). · Czas wstępnego odparowania: ok. 10 - 20 minut. · Czas układania: ok. 15 - 20 minut")
    render_potrzebne_materialy(dane, rep)

# --- SEKCJA: WYKŁADZINA DYWANOWA ---
def generate_report_wykladzina_dywanowa(dane, rep):
    if dane['needs_levelling'] == "NIE" and dane['already_levelled'] == "NIE":
        rep.error("BŁĄD: Pod wykładzinę dywanową wymagane jest wyrównanie podłoża. Poinformuj klienta o konieczności wylania masy!")
        return
        
    render_wspolne_dane_optyczne(dane, rep)
    rep.markdown("#### **II. Zalecenia techniczne (Wykładzina dywanowa)**")
    
    if dane['already_levelled'] == "TAK":
        rep.write("**a) przygotowanie podłoża:**")
        if dane['curing_not_done']:
            rep.write("* **Konieczność przeprowadzenia pełnego procesu wygrzewania podłoża** zgodnie z protokołem.")
            rep.write("Po **przeprowadzeniu pełnego procesu wygrzewania** zalecamy:")
        rep.write("* Szlif podłoża w celu uzyskania gładkiej powierzchni.")
        rep.write("* Dokładne odkurzenie powierzchni odkurzaczem przemysłowym.")
        rep.write("**b) klejenie wykładziny tekstylnej:**")
        rep.write(f"Klejenie wykładziny tekstylnej należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3308']['name']} (szpachla TKB B1 400-450 g/m²). · Czas wstępnego odparowania: ok. 5-10 minut. · Czas otwarty kleju ok. 10-15 minut")
        render_potrzebne_materialy(dane, rep)
        return

    render_wspolne_zalecenia_podloze(dane, rep)
    used_d3004 = render_wspolna_chemia(dane, rep)

    if dane['needs_levelling'] == "TAK" and dane.get('bruzdowane_wybor') != "masa samorozlewna":
        _pu_applied = any(k in dane.get('written_texts', set()) for k in ['PU 280 (1W)', 'PU 280 (Bariera)', 'PU 280 (Bariera Płyta)', 'PU 235 (1W)', 'PU 235 (Bariera)'])
        _skip_d3045 = dane.get('leveling_mesh') == "z siatką" and _pu_applied
        if not used_d3004 and not _skip_d3045:
            if dane.get('firma') == "Mapei":
                write_and_track(dane, rep, 'D 3045')
            else:
                rep.write("* Następnie należy zaaplikować specjalistyczny mostek sczepny za pomocą produktu **WAKOL D 3045**. Aplikować równomiernie za pomocą wałka. Zużycie wynosi **ok. 150 g/m²**. **Czas schnięcia 1 godzina**.")
                write_and_track(dane, rep, 'D 3045')
        if dane.get('leveling_mesh') == "z siatką":
            area = dane.get('area_m2') or 0
            firma_is_mapei = (dane.get('firma') == "Mapei")
            ratio = 8.0 if firma_is_mapei else 7.0
            
            if firma_is_mapei:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} kg na 25 kg masy). Czas schnięcia min. 3h.")
            else:
                rep.write(f"* Na przygotowane podłoże należy rozłożyć matę zbrojeniową **{PRODUCTS['AR 150']['name']}** i zaszpachlować ją masą szpachlową **{PRODUCTS['Z 645']['name']}** z dodatkiem plastyfikatora **{PRODUCTS['D 3060']['name']}** ({ratio} litrów na 25 kg masy). Czas schnięcia min. 3h.")
                
            if area > 0:
                write_and_track(dane, rep, 'AR 150', custom_kg=area)
                kg_z645_mesh = area * 2.0
                write_and_track(dane, rep, 'Z 645 (bruzdowane)', custom_kg=kg_z645_mesh)
                write_and_track(dane, rep, 'D 3060', custom_kg=(kg_z645_mesh / 25.0) * ratio)
        write_and_track(dane, rep, 'Z 675')

    rep.write("* Po wyschnięciu masy samorozlewnej zalecamy szlif podłoża w celu uzyskania gładkiej powierzchni oraz dokładne odkurzenie.")

    rep.write("**c) klejenie wykładziny tekstylnej:**")
    rep.write(f"Klejenie wykładziny tekstylnej należy przeprowadzić przy użyciu kleju {PRODUCTS['D 3308']['name']} (szpachla TKB B1 400-450 g/m²). · Czas wstępnego odparowania: ok. 5-10 minut. · Czas otwarty kleju ok. 10-15 minut")
    render_potrzebne_materialy(dane, rep)

# ==========================================
# EXPORT DO DOCX I PDF
# ==========================================
def _add_docx_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Wyczyść element stopki
    ft_elem = footer._element
    for child in list(ft_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            ft_elem.remove(child)

    # Niebieski pasek jako górna ramka paragrafu
    p_border = OxmlElement('w:p')
    pPr_b = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '24')
    top_b.set(qn('w:space'), '4')
    top_b.set(qn('w:color'), '005293')
    pBdr.append(top_b)
    pPr_b.append(pBdr)
    p_border.append(pPr_b)
    ft_elem.append(p_border)

    # Pomocnicze funkcje do budowania XML
    def make_run(text, bold=False, size_half=16, color=None):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_half))
        rPr.append(sz)
        if color:
            cl = OxmlElement('w:color')
            cl.set(qn('w:val'), color)
            rPr.append(cl)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    # Trzy kolumny przez tabstopy — niezawodne w stopkach Word
    # A4 17cm użytkowe = 9639 twipsów; center=4819, right=9639
    TAB_C = '4819'
    TAB_R = '9639'

    def make_tab_row(left, center, right, is_title=False):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        tabs_el = OxmlElement('w:tabs')
        for val, pos in [('center', TAB_C), ('right', TAB_R)]:
            t = OxmlElement('w:tab')
            t.set(qn('w:val'), val)
            t.set(qn('w:pos'), pos)
            tabs_el.append(t)
        pPr.append(tabs_el)
        p.append(pPr)
        sz = 14 if is_title else 16
        col = '005293' if is_title else None
        bd = is_title
        for i, txt in enumerate([left, center, right]):
            if i > 0:
                r_tab = OxmlElement('w:r')
                r_tab.append(OxmlElement('w:tab'))
                p.append(r_tab)
            if txt:
                p.append(make_run(txt, bold=bd, size_half=sz, color=col))
        return p

    footer_rows = [
        ('ZARZĄD',              'ADRES FIRMY',                       'DANE REJESTROWE',  True),
        ('Stephane Moulin',     'ul. Sławęcińska 16, Macierzysz',    'KRS: 0000163623',  False),
        ('Andreas Taddäus Ziobro', '05-850 Ożarów Mazowiecki',       'NIP: 118-13-89-053', False),
        ('biuro@loba-wakol.pl', 'tel.: +48 22 436 24 20',            'REGON: 013285030', False),
        ('',                    'fax: +48 22 436 24 21',             '',                 False),
    ]
    for left, center, right, is_title in footer_rows:
        ft_elem.append(make_tab_row(left, center, right, is_title))

    ft_elem.append(OxmlElement('w:p'))

def _add_docx_header(doc, data_badania_str='', autor_str=''):
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(4.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Pusty nagłówek Word na wszystkich stronach (brak miejsca na nagłówek)
    hdr = section.header
    hdr_el = hdr._element
    for child in list(hdr_el):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            hdr_el.remove(child)
    hdr_el.append(OxmlElement('w:p'))

    # Logo + dane firmy jako tabela w TREŚCI dokumentu (pierwsza strona naturalnie)
    tbl = doc.add_table(rows=1, cols=2)
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tbl_pr)
    tbl_bdr = OxmlElement('w:tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'none')
        tbl_bdr.append(b)
    tbl_pr.append(tbl_bdr)

    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    # Lewa komórka: logo
    para_logo = left_cell.paragraphs[0]
    para_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists('loba_wakol_logo.png'):
        try:
            para_logo.add_run().add_picture('loba_wakol_logo.png', width=Inches(3.5))
        except:
            pass

    # Prawa komórka: dane firmy
    info_lines = [
        ('Loba-Wakol Polska Sp. z o.o.', True, 14),
        ('ul. Sławęcińska 16, Macierzysz', False, 9),
        ('05-850 Ożarów Mazowiecki', False, 9),
        (f'data: {data_badania_str}', False, 9),
        (f'autor: {autor_str}', False, 9),
        ('tel.: +48 22 436 24 20  |  fax: +48 22 436 24 21', False, 9),
        ('biuro@loba-wakol.pl', False, 9),
    ]
    for i, (text, bold, size) in enumerate(info_lines):
        para = right_cell.paragraphs[0] if i == 0 else right_cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if bold:
            run.font.color.rgb = RGBColor(0x00, 0x52, 0x93)

    # Niebieski separator po tabeli nagłówka
    sep = doc.add_paragraph()
    sep_pPr = sep._p.get_or_add_pPr()
    sep_bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '005293')
    sep_bdr.append(bot)
    sep_pPr.append(sep_bdr)

def compress_image(img_data, max_width=800):
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(img_data))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        if img.width > max_width:
            ratio = max_width / float(img.width)
            new_height = int((float(img.height) * float(ratio)))
            # Użycie LANCZOS dla nowszych wersji PIL, lub ANTIALIAS dla starszych
            try:
                resample = Image.Resampling.LANCZOS
            except AttributeError:
                resample = Image.ANTIALIAS
            img = img.resize((max_width, new_height), resample)
        out_bio = io.BytesIO()
        img.save(out_bio, format="JPEG", quality=75, optimize=True)
        return out_bio.getvalue()
    except ImportError:
        import streamlit as st
        st.warning("Brak biblioteki Pillow (PIL). Zdjęcia nie będą kompresowane.")
        return img_data
    except Exception as e:
        return img_data

def generate_docx(md_text, data_badania_str='', autor_str='', images=None):
    doc = Document()
    # Usuń domyślny pusty paragraf
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Nagłówek (logo + dane firmy) jako treść — PRZED resztą
    _add_docx_header(doc, data_badania_str, autor_str)

    for line in md_text.split('\n\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('#### '):
            p = doc.add_heading(level=2)
            _add_runs(p, line.replace('#### ', ''))
        elif line.startswith('### '):
            p = doc.add_heading(level=1)
            _add_runs(p, line.replace('### ', ''))
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    if images:
        has_images = any(len(img_list) > 0 for img_list in images.values())
        if has_images:
            doc.add_page_break()
            p_attach = doc.add_heading(level=1)
            _add_runs(p_attach, "ZAŁĄCZNIKI ZDJĘCIOWE")
            
            from docx.shared import Inches
            for category, img_list in images.items():
                if img_list:
                    p_cat = doc.add_heading(level=2)
                    _add_runs(p_cat, category)
                    
                    for img in img_list:
                        try:
                            img.seek(0)
                            img_data = compress_image(img.read())
                            p_img = doc.add_paragraph()
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p_img.add_run()
                            run.add_picture(io.BytesIO(img_data), width=Inches(5.0))
                        except Exception as e:
                            doc.add_paragraph(f"[Błąd ładowania zdjęcia: {e}]")

    _add_docx_footer(doc)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _add_runs(p, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 != 0:
            run.bold = True

class WakolPDF(FPDF):
    def __init__(self, data_badania_str, autor_str):
        super().__init__()
        self.data_badania_str = data_badania_str
        self.autor_str = autor_str

    def header(self):
        if self.page_no() == 1:
            try:
                import os
                if os.path.exists('loba_wakol_logo.png'):
                    self.image('loba_wakol_logo.png', x=10, y=8, w=90)
            except:
                pass
            try:
                self.set_font('Arial', 'B', 16)
            except:
                pass
            self.cell(0, 10, 'Loba-Wakol Polska Sp. z o.o.', ln=True, align='R')
            try:
                self.set_font('Arial', '', 9)
            except:
                pass
            label_x = 140
            value_x = 155
            
            def print_row(lbl, val):
                self.set_x(label_x)
                self.cell(15, 4, lbl)
                self.set_x(value_x)
                self.cell(0, 4, val, ln=True)

            print_row("adres:", "Sławęcińska 16, Macierzysz")
            print_row("", "05-850 Ożarów Mazowiecki")
            print_row("data:", self.data_badania_str)
            print_row("autor:", self.autor_str)
            print_row("telefon:", "+48 22 436 24 20")
            print_row("telefax:", "+48 22 436 24 21")
            print_row("e-mail:", "biuro@loba-wakol.pl")
            print_row("strona:", f"{self.page_no()} z {{nb}}")
            self.set_y(60)
        else:
            try:
                self.set_font('Arial', '', 9)
            except:
                pass
            self.cell(0, 10, f"strona {self.page_no()} z {{nb}}", ln=True, align='L')
            self.set_y(20)

    def footer(self):
        footer_y = -32
        try:
            self.set_font('Arial', '', 7)
        except:
            pass

        # Niebieska linia
        self.set_y(footer_y)
        self.set_draw_color(0, 82, 147)
        self.set_line_width(0.8)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)

        col1_x = 10
        col2_x = 75
        col3_x = 150
        row_y = self.get_y()

        # Kolumna 1: ZARZĄD
        try: self.set_font('Arial', 'B', 6)
        except: pass
        self.set_xy(col1_x, row_y); self.cell(60, 3, "ZARZĄD", ln=True)
        try: self.set_font('Arial', '', 7)
        except: pass
        self.set_x(col1_x); self.cell(60, 3, "Stephane Moulin", ln=True)
        self.set_x(col1_x); self.cell(60, 3, "Andreas Taddaeus Ziobro", ln=True)
        self.set_x(col1_x); self.cell(60, 3, "biuro@loba-wakol.pl", ln=True)

        # Kolumna 2: ADRES FIRMY
        try: self.set_font('Arial', 'B', 6)
        except: pass
        self.set_xy(col2_x, row_y); self.cell(70, 3, "ADRES FIRMY", ln=True)
        try: self.set_font('Arial', '', 7)
        except: pass
        self.set_x(col2_x); self.cell(70, 3, "ul. Slawecinska 16, Macierzysz", ln=True)
        self.set_x(col2_x); self.cell(70, 3, "05-850 Ozarow Mazowiecki", ln=True)
        self.set_x(col2_x); self.cell(70, 3, "tel.: +48 22 436 24 20 | fax: +48 22 436 24 21", ln=True)

        # Kolumna 3: DANE REJESTROWE
        try: self.set_font('Arial', 'B', 6)
        except: pass
        self.set_xy(col3_x, row_y); self.cell(50, 3, "DANE REJESTROWE", ln=True)
        try: self.set_font('Arial', '', 7)
        except: pass
        self.set_x(col3_x); self.cell(50, 3, "KRS: 0000163623", ln=True)
        self.set_x(col3_x); self.cell(50, 3, "NIP: 118-13-89-053", ln=True)
        self.set_x(col3_x); self.cell(50, 3, "REGON: 013285030", ln=True)
        
        try:
            import urllib.request
            if not __import__('os').path.exists('wakol_logo.png'):
                urllib.request.urlretrieve('https://www.wakol.com/fileadmin/templates/images/wakol_logo.png', 'wakol_logo.png')
            self.image('wakol_logo.png', x=160, y=-25, w=35)
        except:
            pass

def generate_pdf(md_text, data_badania_str, autor_str, images=None):
    pdf = WakolPDF(data_badania_str, autor_str)
    pdf.alias_nb_pages()
    import os
    try:
        if os.path.exists(r'C:\Windows\Fonts\arial.ttf'):
            pdf.add_font('Arial', '', r'C:\Windows\Fonts\arial.ttf')
            pdf.add_font('Arial', 'B', r'C:\Windows\Fonts\arialbd.ttf')
        elif os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'):
            pdf.add_font('Arial', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf')
            pdf.add_font('Arial', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf')
        else:
            # Download a very reliable font source (Google Fonts - Roboto)
            import urllib.request
            url_reg = 'https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf'
            url_bold = 'https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Bold.ttf'
            if not os.path.exists('Roboto-Regular.ttf'):
                urllib.request.urlretrieve(url_reg, 'Roboto-Regular.ttf')
            if not os.path.exists('Roboto-Bold.ttf'):
                urllib.request.urlretrieve(url_bold, 'Roboto-Bold.ttf')
            pdf.add_font('Arial', '', 'Roboto-Regular.ttf')
            pdf.add_font('Arial', 'B', 'Roboto-Bold.ttf')
        pdf.set_font('Arial', size=11)
    except Exception as e:
        import streamlit as st
        st.error(f"Błąd ładowania czcionki: {str(e)}")
        pdf.set_font('helvetica', size=11)
        
    pdf.set_auto_page_break(auto=True, margin=38)
    pdf.add_page()
    for line in md_text.split('\n\n'):
        line = line.strip()
        if not line: continue
        
        # Proste parsowanie nagłówków
        if line.startswith('#### '):
            pdf.set_font(pdf.font_family, 'B', 14)
            line = line.replace('#### ', '').replace('**', '')
            pdf.multi_cell(0, 8, txt=line)
            pdf.set_font(pdf.font_family, '', 11)
        elif line.startswith('### '):
            pdf.set_font(pdf.font_family, 'B', 16)
            line = line.replace('### ', '').replace('**', '')
            pdf.multi_cell(0, 10, txt=line)
            pdf.set_font(pdf.font_family, '', 11)
        else:
            # Ręczna obsługa pogrubień dla FPDF
            if '**' in line:
                try:
                    pdf.multi_cell(0, 6, txt=line, markdown=True)
                except TypeError:
                    pdf.multi_cell(0, 6, txt=line.replace('**', ''))
            else:
                pdf.multi_cell(0, 6, txt=line)
        pdf.ln(2)
        
    if images:
        has_images = any(len(img_list) > 0 for img_list in images.values())
        if has_images:
            pdf.add_page()
            pdf.set_font(pdf.font_family, 'B', 16)
            pdf.multi_cell(0, 10, txt="ZAŁĄCZNIKI ZDJĘCIOWE")
            pdf.ln(5)
            
            import tempfile
            import os
            for category, img_list in images.items():
                if img_list:
                    pdf.set_font(pdf.font_family, 'B', 14)
                    pdf.multi_cell(0, 8, txt=category)
                    pdf.ln(2)
                    
                    for img in img_list:
                        try:
                            img.seek(0)
                            img_data = compress_image(img.read())
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                                tmp.write(img_data)
                                tmp_path = tmp.name
                                
                            pdf.image(tmp_path, w=130)
                            pdf.ln(5)
                            os.remove(tmp_path)
                        except Exception as e:
                            pdf.set_font(pdf.font_family, '', 11)
                            pdf.multi_cell(0, 6, txt=f"[Błąd ładowania zdjęcia: {e}]")
                            pdf.ln(2)

    output = pdf.output(dest='S')
    if type(output) is str:
        return output.encode('latin-1')
    return bytes(output)

def render_wersja_pro(nazwa_klienta, miejscowosc, adres, autor, data_badania):
    st.markdown("### 🛠️ Wersja PRO - Ręczny Kreator Protokołu")
    
    pro_okladzina = st.text_input("Nazwa okładziny docelowej (do wydruku):", placeholder="np. deska lita, podłoga warstwowa, PCV...")
    pro_area = st.number_input("Powierzchnia inwestycji (m²):", min_value=1.0, step=1.0, format="%.1f", value=None)
    
    st.write("#### 1. Opis protokołu")
    opis_pro = st.text_area("Opis sytuacji / uwagi techniczne:", height=150, placeholder="Wpisz opis stanu podłoża, warunków, uwag technicznych...")

    st.write("#### 2. Czynności przygotowawcze")
    czynnosci = st.text_area("Wpisz czynności (np. Szlifowanie, odkurzanie, nacinanie pęknięć):", height=100)

    st.write("#### 3. Wybór chemii WAKOL")

    if "pro_products_count" not in st.session_state:
        st.session_state.pro_products_count = 1

    product_keys = ["BRAK"] + list(PRODUCTS.keys())
    pro_selected_products = []

    for i in range(st.session_state.pro_products_count):
        st.markdown(f"**Produkt nr {i+1}**")
        col1, col2 = st.columns([2, 1])
        with col1:
            prod_key = st.selectbox(f"Wybierz produkt {i+1}", product_keys, key=f"pro_prod_{i}")

        pro_label = "kg"
        if prod_key != "BRAK":
            if "AR 150" in prod_key or "EM 140" in prod_key:
                pro_label = "m²"
            elif "Płyta RP" in prod_key:
                pro_label = "szt"

        with col2:
            prod_usage = st.number_input(f"Łączna ilość [{pro_label}]", min_value=0.0, step=0.5, format="%.2f", key=f"pro_usage_{i}")

        if prod_key != "BRAK":
            pro_selected_products.append({"key": prod_key, "usage": prod_usage})
            
    if st.button("➕ Dodaj kolejny produkt"):
        st.session_state.pro_products_count += 1
        st.rerun()
        
    st.divider()
    include_cost = st.checkbox("Dołącz wstępny kosztorys materiałowy do protokołu (Netto)", value=True, key="pro_cost")
    
    if st.button("GENERUJ PROTOKÓŁ PRO", type="primary", use_container_width=True):
        if not pro_area:
            st.error("Proszę podać powierzchnię inwestycji (m²)!")
            return
            
        st.divider()
        insert_header()
        rep = ReportBuilder()
        
        # Nagłówek
        tytul = f"Dotyczy: Protokół z robót przygotowawczych w obiekcie:\nAdres: {adres}, {miejscowosc}\nDla: {nazwa_klienta}\n\nSzanowni Państwo,\n\nW dniu {data_badania.strftime('%d.%m.%Y')} zalecono następujący system przygotowania podłoża pod instalację okładziny: {pro_okladzina}.\n\n"
        rep.write(tytul)
        
        if opis_pro.strip():
            rep.write(opis_pro.strip())
            rep.write("\n---\n")

        if czynnosci.strip():
            rep.write("**Sposób przygotowania:**")
            rep.write(f"* {czynnosci.strip()}")

        # Zbieramy materiały do kosztorysu bez wypisywania tekstów produktów
        selected_keys = [p['key'] for p in pro_selected_products if p['usage'] > 0]
        dane_pro = {
            'written_texts': set(selected_keys),  # blokuje wypisywanie opisów produktów
            'materials': [],
            'area_m2': pro_area,
            'include_cost': include_cost,
        }

        for p in pro_selected_products:
            key = p['key']
            needed = p['usage']
            if needed > 0:
                write_and_track(dane_pro, rep, key, custom_kg=needed)
            else:
                st.warning(f"Produkt {PRODUCTS[key]['name']} ma ustawione zużycie 0. Zostanie pominięty w kosztorysie.")
                
        rep.write("\n---\n")
        render_potrzebne_materialy(dane_pro, rep)
        
        rep.write("\n**Prosimy o zapoznanie się z kartami technicznymi zalecanych produktów WAKOL.**\n\nPodstawą naszego zalecenia jest stosowanie i prawidłowa obróbka wszystkich wymienionych materiałów firmy WAKOL w podanej kolejności, przestrzegając reguł rzemiosła i obowiązujących norm oraz instrukcji.\n\nW przypadku jakichkolwiek pytań lub wątpliwości proszę o kontakt pod numer telefonu: 603 214 218\n\nZ poważaniem,\n\nLoba-Wakol Polska Sp. z o.o.\n" + autor)
        
        st.markdown(rep.get_markdown())
        
        if EXPORTS_READY:
            col_d1, col_d2 = st.columns(2)
            safe_adres = adres.replace(' ', '_').replace('/', '_').replace('.', '')
            data_str = data_badania.strftime('%d-%m-%Y')
            safe_klient = nazwa_klienta.replace(' ', '_').replace('/', '_')
            base_filename = f"Protokol_PRO_Wakol_{safe_klient}_{safe_adres}_{data_str}"
            
            with col_d1:
                docx_file = generate_docx(rep.get_markdown(), data_badania.strftime('%d.%m.%Y'), autor, None)
                st.download_button(label="📄 Pobierz jako plik Word (.docx)", data=docx_file, file_name=f"{base_filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            with col_d2:
                pdf_file = generate_pdf(rep.get_markdown(), data_badania.strftime('%d.%m.%Y'), autor, None)
                if pdf_file:
                    st.download_button(label="📕 Pobierz jako plik PDF (.pdf)", data=pdf_file, file_name=f"{base_filename}.pdf", mime="application/pdf", use_container_width=True)


# ==========================================
# 3. INTERFEJS UŻYTKOWNIKA (FORMULARZ)
# ==========================================

st.title("📄 Generator Protokołu Oględzin")
with st.container():
    col1, col2 = st.columns(2)
    with col1:
        nazwa_klienta = st.text_input("Nazwa Klienta", "Jan Kowalski")
        miejscowosc = st.text_input("Miejscowość", "Huta Dłutowska")
        adres = st.text_input("Ulica i nr", "ul. Pabianicka 15")
    with col2:
        autor = st.text_input("Autor protokołu", "Przemysław Tyszko")
        data_badania = st.date_input("Data badania", date.today())

st.divider()
tryb = st.radio("Wybierz tryb pracy aplikacji:", ["Wersja Ogólna (Automatyczna)", "Wersja PRO (Ręczna)"], horizontal=True)
st.divider()

if tryb == "Wersja PRO (Ręczna)":
    render_wersja_pro(nazwa_klienta, miejscowosc, adres, autor, data_badania)
    st.stop()


st.markdown("---")
firma = st.radio("🏢 Wybierz firmę:", ["Wakol", "Mapei", "Uzin"], horizontal=True)
if firma == "Mapei":
    PRODUCTS = PRODUCTS_MAPEI
elif firma == "Uzin":
    PRODUCTS = PRODUCTS_UZIN
else:
    PRODUCTS = PRODUCTS_WAKOL
st.markdown("---")

flooring_type = st.selectbox("Wybierz rodzaj okładziny (Sekcja):", ["deska warstwowa", "podłoga laminowana", "lity parkiet (maks. 8 cm x 60 cm)", "mozaika drewniana (min. 16 mm grubości, maks. 20 cm długości)", "deska lita", "wykładzina dywanowa", "pcv w rolce", "lvt cienkie", "lvt grube z twardym rdzeniem"])

klej_typ = None
lvt_bottom_type = None
if flooring_type in ["deska warstwowa", "podłoga laminowana", "lity parkiet (maks. 8 cm x 60 cm)", "mozaika drewniana (min. 16 mm grubości, maks. 20 cm długości)", "deska lita"]:
    klej_typ = st.radio("Rodzaj kleju:", ["elastyczny", "bezprzesuwny"], horizontal=True)
elif flooring_type == "lvt grube z twardym rdzeniem":
    lvt_bottom_type = st.radio("Rodzaj spodu LVT:", ["Winyl z zintegrowanym spodem korkowym", "Winyl na homogenicznym spodzie", "Winyl na piankowym spodzie"], horizontal=True)

st.markdown(f"### Wywiad Techniczny dla: **{flooring_type.upper()}**")

substrate = st.selectbox("1. Rodzaj podłoża", ["jastrych cementowy", "jastrych anhydrytowy", "płyta fundamentowa", "podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB", "płytki ceramiczne", "masa samorozlewna"])

_force_levelling = False
_force_holes = False
has_adhesive_residues = False

if substrate in ["podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB", "płytki ceramiczne"]:
    st.write("1a. Czy podłoże jest stabilnie związane z podkładem?")
    substrate_stable = st.radio("Stabilność podłoża:", ["TAK", "NIE"], index=0, horizontal=True, key="substrate_stable")
    if substrate_stable == "NIE":
        if substrate in ["podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB"]:
            st.warning("⚠️ Podłoże wymaga demontażu. Po demontażu konieczne jest wyrównanie masą samorozlewną.")
            substrate = st.selectbox("Rodzaj podłoża po demontażu:", ["jastrych cementowy", "jastrych anhydrytowy", "płyta fundamentowa"], key="sub_after_demo")
            _force_levelling = True
        else:
            st.warning("⚠️ Podłoże ceramiczne nie jest stabilnie związane z podkładem.")
            ceramic_action = st.radio(
                "Wybierz sposób postępowania:",
                ["Konieczność skucia całości", "Skucie luźnych fragmentów i zaszpachlowanie ubytków"],
                key="ceramic_action"
            )
            if ceramic_action == "Konieczność skucia całości":
                st.warning("⚠️ Po skuciu całości płytek konieczne jest wyrównanie masą samorozlewną.")
                substrate = st.selectbox("Rodzaj podłoża po skuciu:", ["jastrych cementowy", "jastrych anhydrytowy", "płyta fundamentowa"], key="sub_after_demo")
                _force_levelling = True
            else:
                st.info("ℹ️ Skucie luźnych fragmentów ceramicznych i zaszpachlowanie powstałych ubytków.")
                _force_holes = True

area_m2 = st.number_input("Powierzchnia inwestycji (m²):", min_value=1.0, step=1.0, format="%.1f", value=None)
substrate_age_val = st.number_input("Wiek podłoża (podaj ilość miesięcy):", min_value=0.5, step=0.5, format="%.1f", value=None)

st.write("2. Czy jest instalacja ogrzewania podłogowego?")
heating_exists = st.radio("Ogrzewanie:", ["TAK", "NIE"], index=1, horizontal=True)
heating_info = ""; heating_curing_done = None; h_type = None; bruzdowane_wybor = None
if heating_exists == "TAK":
    h_type = st.selectbox("Typ ogrzewania:", ["wodne klasyczne", "bruzdowane", "w suchej zabudowie", "elektryczne (powierzchniowe)", "elektryczne (głębokie)", "płyta fundamentowa grzewcza"])
    if h_type == "bruzdowane":
        if flooring_type == "lvt cienkie":
            bruzdowane_wybor = "masa samorozlewna"
            st.info("Przy LVT cienkim i ogrzewaniu bruzdowanym wymagana jest masa samorozlewna — technologia ustawiona automatycznie.")
        else:
            bruzdowane_wybor = st.radio("Wybierz technologię (ogrzewanie bruzdowane):", ["masa samorozlewna", "płyta RP"], horizontal=True)
        
    if h_type != "bruzdowane":
        st.write("❓ Czy został przeprowadzony proces wygrzewania zgodnie z protokołem?")
        heating_curing_done = st.radio("Proces wygrzewania:", ["TAK", "NIE"], index=1, horizontal=True)
    else:
        heating_curing_done = "TAK"
    mapping = {"wodne klasyczne": "instalacja ogrzewania podłogowego wodna, klasyczna", "bruzdowane": "instalacja ogrzewania podłogowego wodna, bruzdowana", "w suchej zabudowie": "instalacja ogrzewania podłogowego wodna, w suchej zabudowie", "elektryczne (powierzchniowe)": "instalacja ogrzewania podłogowego elektryczna, powierzchniowa", "elektryczne (głębokie)": "instalacja ogrzewania podłogowego elektryczna, umieszczona głęboko w podłożu", "płyta fundamentowa grzewcza": "ogrzewanie realizowane poprzez płytę fundamentową grzewczą"}
    heating_info = mapping.get(h_type, h_type)

# --- LOGIKA NORM I BARIER ---
if substrate == "płyta fundamentowa":
    limit = 1.5 if heating_exists == "TAK" else 1.8
    barrier_max = 2.8
elif substrate == "jastrych anhydrytowy":
    limit = 0.3 if heating_exists == "TAK" else 0.5
    barrier_max = 2.5 if heating_exists == "TAK" else 3.5
else:
    limit = 1.5 if heating_exists == "TAK" else 1.8
    barrier_max = 2.5 if heating_exists == "TAK" else 3.5

# --- WILGOTNOŚĆ PODŁOŻA + DECYZJA (zaraz po ogrzewaniu) ---
_substrate_no_moisture = substrate in ["podłoże drewniane (parkiet, deska)", "podłoże z płyty OSB", "płytki ceramiczne"]
if _substrate_no_moisture:
    st.info("3. Poziom wilgoci podłoża — nie dotyczy tego rodzaju podłoża.")
    moisture = None
elif substrate == "płyta fundamentowa":
    moisture = st.number_input("3. Poziom wilgoci podłoża (%)", format="%.1f", value=None)
else:
    moisture = st.number_input("3. Poziom wilgoci podłoża (CM %)", format="%.1f", value=None)

decision_after_cure = None
needs_drying_action = False

if moisture is not None and moisture > limit:
    needs_drying_action = True
    opt_dry = "przeprowadzenie procesu wygrzewania" if (heating_exists == "TAK" and heating_curing_done == "NIE") else "dalsze osuszanie"
    if substrate == "płyta fundamentowa":
        if moisture > barrier_max:
            st.warning(f"Podłoże jest zbyt wilgotne. Konieczność doprowadzenia do poziomu wilgoci max. {barrier_max}% przed wykonaniem bariery przeciwwilgociowej.")
            decision_after_cure = "osuszanie przed barierą"
        else:
            decision_after_cure = "Wykonanie bariery przeciwwilgociowej"
            needs_drying_action = False
    elif h_type == "bruzdowane":
        st.warning(f"Podłoże jest zbyt wilgotne. Konieczność doprowadzenia do normatywnego poziomu wilgoci ({limit}% CM) przed przystąpieniem do dalszych prac.")
        decision_after_cure = "dalsze osuszanie"
    elif substrate == "jastrych anhydrytowy":
        st.info(f"Dla jastrychu anhydrytowego nie ma możliwości wykonania bariery przeciwwilgociowej. Konieczność doprowadzenia do normatywnego poziomu wilgoci ({limit}% CM).")
        if heating_exists == "TAK" and heating_curing_done == "TAK":
            decision_after_cure = "kolejny proces wygrzewania"
        else:
            decision_after_cure = opt_dry
    else:
        if moisture <= barrier_max:
            decision_after_cure = st.radio("Postępowanie z podwyższoną wilgocią:", ["Wykonanie bariery przeciwwilgociowej", opt_dry], horizontal=True)
            needs_drying_action = (decision_after_cure != "Wykonanie bariery przeciwwilgociowej")
        else:
            decision_after_cure = opt_dry
else:
    if substrate == "płyta fundamentowa":
        decision_after_cure = "Wykonanie bariery przeciwwilgociowej"

st.write("4. Czy są ubytki bądź zdegradowane fragmenty wymagające wypełnienia masą naprawczą?")
hole_details = ""
holes_depth = None
h_depth = None; h_width = None; h_length = None
img_holes = []
if _force_holes:
    holes = "TAK"
    st.info("ℹ️ Ubytki po skuciu płytek ceramicznych — podaj wymiary:")
    col_h1, col_h2, col_h3 = st.columns(3)
    with col_h1: h_depth = st.number_input("Grubość ubytków (cm)", min_value=0.1, value=None, key="h_depth_f")
    with col_h2: h_width = st.number_input("Szerokość ubytków (cm)", min_value=0.1, value=None, key="h_width_f")
    with col_h3: h_length = st.number_input("Długość ubytków (m)", min_value=0.01, value=None, key="h_length_f")
    if h_depth and h_width and h_length: hole_details = f" o wymiarach ok. {h_length} m x {h_width} cm i grubości {h_depth} cm"
    holes_depth = h_depth
    img_holes = st.file_uploader("Zdjęcia ubytków po skuciu:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_holes_f")
else:
    holes = st.radio("Ubytki:", ["TAK", "NIE"], index=1, horizontal=True)
    if holes == "TAK":
        col_h1, col_h2, col_h3 = st.columns(3)
        with col_h1: h_depth = st.number_input("Grubość (cm)", min_value=0.1, value=None)
        with col_h2: h_width = st.number_input("Szerokość (cm)", min_value=0.1, value=None)
        with col_h3: h_length = st.number_input("Długość (m)", min_value=0.01, value=None)
        if h_depth and h_width and h_length: hole_details = f" o wymiarach ok. {h_length} m x {h_width} cm i grubości {h_depth} cm"
        holes_depth = h_depth
        img_holes = st.file_uploader("Zdjęcia ubytków:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_holes")

st.write("3b. Czy wymagane jest miejscowe wyrównanie masą szpachlową?")
local_leveling = st.radio("Miejscowe wyrównanie:", ["TAK", "NIE"], index=1, horizontal=True, key="local_lev")
local_leveling_kg = None
local_leveling_details = ""
if local_leveling == "TAK":
    local_leveling_kg = st.number_input("Ilość masy szpachlowej Z 645 (kg):", min_value=0.1, value=None, key="ll_kg")
    if local_leveling_kg:
        bags_ll = math.ceil(local_leveling_kg / 25)
        st.info(f"Ilość worków Z 645: **{bags_ll} szt. po 25 kg**")

if flooring_type == "lvt cienkie":
    st.write("4. Podłoże zostało wyrównane masą samorozlewną")
else:
    st.write("4. Czy całość podłoża wymaga wyrównania masą samorozlewną?")

leveling_thickness = 0
already_levelled = "NIE"
leveling_mesh = "bez siatki"

if _force_levelling:
    st.info("Wyrównanie jest wymuszone po demontażu podłoża.")
    needs_levelling = "TAK"
    leveling_thickness = st.number_input("Planowana grubość masy po demontażu (mm):", min_value=1, value=None, key="lev_thick_demo")
    leveling_mesh = st.radio("Rodzaj wyrównania:", ["bez siatki", "z siatką"], index=0, horizontal=True, key="lev_mesh_demo")
    st.write("Czy na podłożu są pozostałości starych spoin klejowych?")
    has_adhesive_residues = st.radio("Pozostałości kleju:", ["TAK", "NIE"], index=1, horizontal=True, key="adhesive_res") == "TAK"
elif h_type == "bruzdowane" and bruzdowane_wybor == "masa samorozlewna":
    st.info("Wyrównanie jest wymuszone przez technologię 'masa samorozlewna' na ogrzewaniu bruzdowanym.")
    needs_levelling = "TAK"
    leveling_thickness = 5
    st.info("Grubość masy została automatycznie ustalona na 5 mm.")
elif flooring_type == "lvt cienkie":
    lvt_already = st.radio("Odpowiedź:", ["TAK", "NIE"], index=1, horizontal=True)
    if lvt_already == "TAK":
        already_levelled = "TAK"
        needs_levelling = "NIE"
    else:
        already_levelled = "NIE"
        needs_levelling = "TAK"
        st.info("Wyrównanie podłoża (masą WAKOL Z 675) jest technologicznie wymuszone pod okładzinę LVT cienkie.")
        leveling_thickness = st.number_input("Planowana grubość masy (mm):", min_value=1, value=None)
        leveling_mesh = st.radio("Rodzaj wyrównania:", ["bez siatki", "z siatką"], index=0, horizontal=True)
else:
    needs_levelling = st.radio("Wymaga wyrównania:", ["TAK", "NIE"], index=1, horizontal=True)
    if needs_levelling == "TAK":
        leveling_thickness = st.number_input("Planowana grubość masy (mm):", min_value=1, value=None)
        leveling_mesh = st.radio("Rodzaj wyrównania:", ["bez siatki", "z siatką"], index=0, horizontal=True)
    elif flooring_type in ["wykładzina dywanowa", "pcv w rolce"]:
        st.warning("Pod wybraną okładzinę wymagane jest wyrównanie podłoża.")
        already_levelled = st.radio("Czy podłoże zostało już wcześniej wyrównane?", ["TAK", "NIE"], index=1, horizontal=True)

st.write("5. Czy dylatacje obwodowe zachowane prawidłowo?")
dilatations_obw_ok = st.radio("Dylatacje obwodowe:", ["TAK", "NIE"], index=0, horizontal=True)
st.write("6. Czy występują klawiszujące dylatacje pozorne?")
cracks_klaw = st.radio("Klawiszowanie pozorne:", ["TAK", "NIE"], index=1, horizontal=True)
klaw_meters = 0.0
img_klaw = []
if cracks_klaw == "TAK":
    klaw_meters = st.number_input("Ilość mb klawiszujących:", min_value=0.1, step=0.1, value=None)
    img_klaw = st.file_uploader("Zdjęcia klawiszujących dylatacji:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_klaw")
st.write("7. Czy występują pęknięcia podłoża wymagające zespolenia?")
cracks_pek = st.radio("Pęknięcia do zespolenia:", ["TAK", "NIE"], index=1, horizontal=True)
pek_meters = 0.0
img_pek = []
if cracks_pek == "TAK":
    pek_meters = st.number_input("Ilość mb pęknięć do zespolenia:", min_value=0.1, step=0.1, value=None)
    img_pek = st.file_uploader("Zdjęcia pęknięć:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_pek")

st.write("8. Rodzaj wentylacji")
ventilation_type = st.radio("Wentylacja:", ["Grawitacyjna", "Mechaniczna"], horizontal=True)

dodatkowe_informacje = st.text_area("Dodatkowe informacje z oględzin (opcjonalnie):", placeholder="Wpisz inne zaobserwowane uwagi do protokołu...")
img_dodatkowe = st.file_uploader("Zdjęcia - dodatkowe informacje:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_dod")

col_w1, col_w2 = st.columns(2)
with col_w1: temp_air = st.number_input("9. Temperatura powietrza (°C)", step=0.5, value=None)
with col_w2: hum_air = st.number_input("10. Wilgotność powietrza (%)", step=1.0, value=None)

st.write("### 11. Testy mechaniczne i Wytrzymałość")
col_t1, col_t2, col_t3 = st.columns(3)
with col_t1: test_hammer = st.selectbox("Młotek", ["", "negatywny", "dostateczny", "pozytywny"], index=0)
with col_t2: test_ripper = st.selectbox("Rysik", ["", "negatywny", "dostateczny", "pozytywny"], index=0)
with col_t3: test_brush = st.selectbox("Szczotka", ["", "negatywny", "dostateczny", "pozytywny"], index=0)
img_mech = st.file_uploader("Zdjęcia z testów mechanicznych:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_mech")
st.write("**Badanie PressoMess**")
presso_results = []
_presso_cols = st.columns(6)
for i in range(6):
    with _presso_cols[i]:
        presso_results.append(st.number_input(f"Próba {i+1} (N/mm²)", min_value=0.0, step=0.1, format="%.2f", key=f"p_{i}", value=None))
img_presso = st.file_uploader("Zdjęcia - PressoMess:", accept_multiple_files=True, type=["png", "jpg", "jpeg"], key="img_presso")
strength_labels = {1: "bardzo słaby", 2: "słaby", 3: "umiarkowanie słaby", 4: "umiarkowanie mocny", 5: "mocny"}
if substrate == "podłoże z płyty OSB":
    strength_val = 5
    st.info("Ocena ogólna wytrzymałości podłoża: **mocny** — wartość ustawiona automatycznie dla płyty OSB.")
else:
    strength_val = st.select_slider("Ocena ogólna wytrzymałości podłoża:", options=[1, 2, 3, 4, 5], value=3, format_func=lambda x: strength_labels[x])

st.write("### 13. Opcje raportu")
include_cost = st.checkbox("Dołącz wstępny kosztorys materiałowy do protokołu (Netto)", value=True)

# PAKOWANIE DANYCH DO SŁOWNIKA DLA FUNKCJI GENERUJĄCYCH
dane_protokolu = {
    "include_cost": include_cost,
    "firma": firma,
    "flooring_type": flooring_type,
    "substrate": substrate,
    "area_m2": area_m2,
    "klej_typ": klej_typ,
    "lvt_bottom_type": lvt_bottom_type,
    "substrate_age_val": substrate_age_val,
    "barrier_max": barrier_max,
    "heating_exists": heating_exists,
    "heating_info": heating_info,
    "heating_curing_done": heating_curing_done,
    "h_type": h_type,
    "bruzdowane_wybor": bruzdowane_wybor,
    "needs_levelling": needs_levelling,
    "leveling_thickness": leveling_thickness,
    "leveling_mesh": leveling_mesh,
    "already_levelled": already_levelled,
    "requires_demolition": _force_levelling,
    "has_adhesive_residues": has_adhesive_residues,
    "local_leveling": local_leveling,
    "local_leveling_kg": local_leveling_kg,
    "local_leveling_details": local_leveling_details,
    "dilatations_obw_ok": dilatations_obw_ok,
    "cracks_klaw": cracks_klaw,
    "klaw_meters": klaw_meters,
    "cracks_pek": cracks_pek,
    "pek_meters": pek_meters,
    "holes": holes,
    "holes_depth": holes_depth if 'holes_depth' in locals() else None,
    "holes_width": h_width if 'h_width' in locals() else None,
    "holes_length": h_length if 'h_length' in locals() else None,
    "hole_details": hole_details,
    "ventilation_type": ventilation_type,
    "dodatkowe_informacje": dodatkowe_informacje,
    "moisture": moisture,
    "limit": limit,
    "curing_not_done": (heating_exists == "TAK" and heating_curing_done == "NIE"),
    "is_moisture_neg": (moisture is not None and moisture > limit),
    "norm_val_bracket": f"({limit}% CM)",
    "decision_after_cure": decision_after_cure,
    "needs_drying_action": needs_drying_action,
    "test_hammer": test_hammer,
    "test_ripper": test_ripper,
    "test_brush": test_brush,
    "strength_labels": strength_labels,
    "strength_val": strength_val,
    "temp_air": temp_air,
    "hum_air": hum_air,
    "presso_results": presso_results,
    "images": {
        "Klawiszujące dylatacje": img_klaw if 'img_klaw' in locals() and img_klaw else [],
        "Pęknięcia podłoża": img_pek if 'img_pek' in locals() and img_pek else [],
        "Ubytki w podłożu": img_holes if 'img_holes' in locals() and img_holes else [],
        "Testy mechaniczne (rysik, młotek, szczotka)": img_mech if 'img_mech' in locals() and img_mech else [],
        "Badanie PressoMess": img_presso if 'img_presso' in locals() and img_presso else [],
        "Dodatkowe informacje z oględzin": img_dodatkowe if 'img_dodatkowe' in locals() and img_dodatkowe else []
    }
}

# --- GENEROWANIE PROTOKOŁU W ZALEŻNOŚCI OD WYBRANEJ OKŁADZINY ---
if st.button(f"GENERUJ PROTOKÓŁ OGLĘDZIN DLA: {flooring_type.upper()}", type="primary", use_container_width=True):
    if moisture is None and not _substrate_no_moisture:
        st.error("Proszę podać wilgotność podłoża!")
    else:
        st.divider()
        insert_header()
        
        rep = ReportBuilder()
        
        # Generowanie nagłówka do DOC/PDF
        tytul = f"Dotyczy: Protokół z oględzin inwestycji w obiekcie:\nAdres: {adres}, {miejscowosc}\nDla: {nazwa_klienta}\n\nSzanowni Państwo,\n\nW dniu {data_badania.strftime('%d.%m.%Y')} dokonano wstępnych oględzin i pomiarów wytrzymałości podłoża ({substrate}) oraz pomiaru wilgotności przed przyklejeniem okładziny ({flooring_type}).\n\n"
        rep.write(tytul)
        
        rep.markdown("#### **I. Oględziny i badania**")
        
        if flooring_type in ["deska warstwowa", "podłoga laminowana", "lity parkiet (maks. 8 cm x 60 cm)", "mozaika drewniana (min. 16 mm grubości, maks. 20 cm długości)"]:
            generate_report_deska_warstwowa(dane_protokolu, rep)
        elif flooring_type == "deska lita":
            generate_report_deska_lita(dane_protokolu, rep)
        elif flooring_type == "lvt cienkie":
            generate_report_lvt_cienkie(dane_protokolu, rep)
        elif flooring_type == "pcv w rolce":
            generate_report_pcv_w_rolce(dane_protokolu, rep)
        elif flooring_type == "wykładzina dywanowa":
            generate_report_wykladzina_dywanowa(dane_protokolu, rep)
        elif flooring_type == "lvt grube z twardym rdzeniem":
            generate_report_lvt_grube(dane_protokolu, rep)
        else:
            rep.error("Nieobsługiwany typ okładziny.")
            
        rep.write("\n**Prosimy o zapoznanie się z kartami technicznymi zalecanych produktów WAKOL.**\n\nPodstawą naszego zalecenia jest stosowanie i prawidłowa obróbka wszystkich wymienionych materiałów firmy WAKOL w podanej kolejności, przestrzegając reguł rzemiosła i obowiązujących norm oraz instrukcji.\n\nW przypadku jakichkolwiek pytań lub wątpliwości proszę o kontakt pod numer telefonu: 603 214 218\n\nZ poważaniem,\n\nLoba-Wakol Polska Sp. z o.o.\n" + autor)
        
        # Wyświetlenie na ekranie (cel użytkownika)
        st.markdown(rep.get_markdown())

        st.markdown("""
<div style="font-family: Arial, sans-serif; color: #333; max-width: 800px;">
    <div style="width: 100%; height: 4px; background-color: #005293; margin-bottom: 20px;"></div>
    <table style="width: 100%; border-collapse: collapse; margin-bottom: 25px;">
        <tr>
            <td style="width: 20%; text-align: left; vertical-align: middle;">
                <img src="https://www.loba-wakol.pl/fileadmin/templates/images/loba_logo.png" alt="Loba" style="height: 50px; width: auto; display: block;">
            </td>
            <td style="width: 60%; text-align: center; vertical-align: middle;">
                <h1 style="font-size: 12.5pt; margin: 0; color: #005293; text-transform: uppercase; white-space: nowrap; letter-spacing: 0.2px;">
                    LOBA-WAKOL POLSKA SPÓŁKA Z O.O.
                </h1>
            </td>
            <td style="width: 20%; text-align: right; vertical-align: middle;">
                <img src="https://www.loba-wakol.pl/fileadmin/templates/images/wakol_logo.png" alt="Wakol" style="height: 50px; width: auto; display: block; margin-left: auto;">
            </td>
        </tr>
    </table>
    <div style="border-top: 1px solid #005293; padding-top: 15px;">
        <table style="width: 100%; font-size: 9pt; line-height: 1.5; border-collapse: collapse;">
            <tr>
                <td style="width: 33%; vertical-align: top;">
                    <span style="color: #005293; font-size: 8pt; font-weight: bold; display: block; margin-bottom: 4px; text-transform: uppercase;">ZARZĄD</span>
                    <strong>Stephane Moulin</strong><br>
                    <strong>Andreas Taddäus Ziobro</strong><br>
                    <a href="mailto:biuro@loba-wakol.pl" style="color: #005293; text-decoration: none;">biuro@loba-wakol.pl</a>
                </td>
                <td style="width: 34%; vertical-align: top; text-align: center;">
                    <span style="color: #005293; font-size: 8pt; font-weight: bold; display: block; margin-bottom: 4px; text-transform: uppercase;">ADRES FIRMY</span>
                    ul. Sławęcińska 16, Macierzysz<br>
                    05-850 Ożarów Mazowiecki<br>
                    tel.: +48 22 436 24 20<br>
                    fax: +48 22 436 24 21
                </td>
                <td style="width: 33%; vertical-align: top; text-align: right;">
                    <span style="color: #005293; font-size: 8pt; font-weight: bold; display: block; margin-bottom: 4px; text-transform: uppercase;">DANE REJESTROWE</span>
                    KRS: 0000163623<br>
                    NIP: 118-13-89-053<br>
                    REGON: 013285030
                </td>
            </tr>
        </table>
    </div>
</div>
""", unsafe_allow_html=True)

        st.divider()
        
        # Przyciski pobierania
        if EXPORTS_READY:
            col_d1, col_d2 = st.columns(2)
            
            safe_adres = adres.replace(' ', '_').replace('/', '_').replace('.', '')
            data_str = data_badania.strftime('%d-%m-%Y')
            safe_klient = nazwa_klienta.replace(' ', '_').replace('/', '_')
            base_filename = f"Protokol_Wakol_{safe_klient}_{safe_adres}_{data_str}"
            
            with col_d1:
                docx_file = generate_docx(rep.get_markdown(), data_badania.strftime('%d.%m.%Y'), autor, dane_protokolu.get('images'))
                st.download_button(
                    label="📄 Pobierz jako plik Word (.docx)",
                    data=docx_file,
                    file_name=f"{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col_d2:
                pdf_file = generate_pdf(rep.get_markdown(), data_badania.strftime('%d.%m.%Y'), autor, dane_protokolu.get('images'))
                if pdf_file:
                    st.download_button(
                        label="📕 Pobierz jako plik PDF (.pdf)",
                        data=pdf_file,
                        file_name=f"{base_filename}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
        else:
            st.error("⚠️ Brak bibliotek do generowania Word/PDF. Dodaj plik `requirements.txt` w swoim repozytorium na GitHubie z zawartością:\n```\npython-docx\nfpdf2\n```")
